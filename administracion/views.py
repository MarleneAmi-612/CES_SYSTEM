import io
import os
import json
import uuid
import qrcode
import fitz           
import pikepdf
import zipfile
import secrets
import logging
import inspect
import tempfile
import subprocess
import unicodedata
import requests
from io import BytesIO
from docx import Document
from base64 import b64encode
from datetime import timedelta,datetime,date
from pathlib import Path
from typing import Optional   
from alumnos.models import Request
from types import SimpleNamespace
from urllib.request import Request as UrlReq, urlopen
from django.contrib.staticfiles.finders import find as find_static
from email.mime.image import MIMEImage
from django.core.mail import EmailMultiAlternatives
from django.utils.timezone import localdate
from docxtpl import DocxTemplate

# ============================
# Terceros (third-party)
# ============================
import requests               
from PIL import Image                 
from weasyprint import HTML
from csp import constants as csp
from csp.decorators import csp_update
from django_otp.decorators import otp_required

# ============================
# Django
# ============================
from django.apps import apps
from django.conf import settings
from django.contrib import messages
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth import (
    authenticate,
    login,
    logout,
    get_user_model,
    update_session_auth_hash,
)
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib.auth.models import Group
from django.contrib.auth.views import (
    PasswordResetView,
    PasswordResetDoneView,
    PasswordResetConfirmView,
    PasswordResetCompleteView,
)
from django.core.exceptions import PermissionDenied
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.core.paginator import Paginator
from django.core.mail import send_mail,EmailMessage, EmailMultiAlternatives
from django.core.mail import EmailMessage
from django.db import connections, transaction,IntegrityError
from django.db.models import Q, Count, Max
from django.forms import modelform_factory
from django.http import (
    JsonResponse,
    Http404,
    HttpResponse,
    FileResponse,
    HttpResponseBadRequest,
    HttpResponseNotAllowed,
    HttpResponseForbidden,
    HttpResponseRedirect,
    HttpResponseNotFound,
)
from django.middleware.csrf import get_token
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.urls import reverse, reverse_lazy
from django.utils import timezone
from django.utils.decorators import method_decorator
from django.utils.html import escape
from django.utils.dateparse import parse_date
from django.utils.text import slugify
from django.views.decorators.csrf import csrf_exempt, csrf_protect
from django.views.decorators.http import (
    require_GET,
    require_POST,
    require_http_methods,
)
from django.views.generic import TemplateView, FormView

# ============================
# Apps locales
# ============================
from alumnos.models import Request  # u otros modelos que sí existan
from administracion.models import Graduate,DocToken,DiplomaBackground,RejectedArchive,Program,ProgramSim,DesignTemplate,ConstanciaType,DocxTemplate
egresados = Graduate.objects.all()

# App "alumnos"
from alumnos.models import (
    Request,
    RequestEvent,
    Program as ProgramSim,
)

# App "administracion"
from .forms import AdminLoginForm
from .models import (
    Graduate,
    AdminAccessLog,
    DocToken,
    DesignTemplate,
    DesignTemplateVersion,
    TemplateAsset,
    ProgramSim,
    CertificateType,
    Program,              
    Program as ProgramAdmin,  
    ConstanciaType,
    DocxTemplate,          
)
from .security import failed_attempts_count
from .thumbs import save_thumb
from .models import Program as ProgramAdmin, ConstanciaType
from django.templatetags.static import static
from django.contrib.staticfiles import finders
from docx import Document  
# ============================
# Config / helpers globales
# ============================
User = get_user_model()

# Si estás en Windows y quieres rasterizar con Poppler / LibreOffice
POPPLER_BIN = os.environ.get("POPPLER_BIN")
SOFFICE_BIN = os.environ.get("SOFFICE_BIN")

# Egresado puede no existir en todos los entornos
try:
    from .models import Egresado
except Exception:
    Egresado = None

log = logging.getLogger(__name__)
# Si ya tienes una función parecida para DC3, puedes reutilizarla o adaptarla.
MESES_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]
# =========================
#  CONTEXTO PARA PLANTILLAS DOCX
# =========================

def _fecha_es_letra(fecha):
    """
    Convierte una fecha date a '11 de noviembre de 2025'.
    Si viene None, regresa cadena vacía.
    """
    if not fecha:
        return ""
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    return f"{fecha.day} de {meses[fecha.month - 1]} de {fecha.year}"


def _docx_context_for_request(req, tipo_real: str, *, token: str = "", verify_url: str = ""):
    """
    Construye el diccionario de variables que van a usar las plantillas DOCX.

    Se usa tanto para diplomas como para constancias (dc3 / cproem).
    Aquí NO se hace nada de DOCX todavía, solo preparamos datos.
    """

    # Nombre separado
    nombre = (getattr(req, "name", "") or "").strip()
    apellido = (getattr(req, "lastname", "") or "").strip()
    full_name = f"{nombre} {apellido}".strip()

    # Programa
    program_obj = getattr(req, "program", None)
    programa = getattr(program_obj, "name", "") or ""

    # Graduate (si existe)
    grad = getattr(req, "graduate", None)

    # Fechas de validez / curso
    if grad and (getattr(grad, "validity_start", None) or getattr(grad, "validity_end", None)):
        inicio = getattr(grad, "validity_start", None)
        fin = getattr(grad, "validity_end", None)
    else:
        inicio = getattr(req, "start_date", None)
        fin = getattr(req, "end_date", None)

    # Horas (intensidad)
    # Prioridad: Graduate.hours -> Request.hours -> None
    horas = None
    if grad and hasattr(grad, "hours") and grad.hours is not None:
        horas = grad.hours
    elif hasattr(req, "hours") and req.hours is not None:
        horas = req.hours

    # CURP / RFC / etc.
    curp = getattr(req, "curp", "") or ""
    rfc = getattr(req, "rfc", "") or ""
    puesto = getattr(req, "job_title", "") or ""
    giro = getattr(req, "industry", "") or ""
    razon_social = getattr(req, "business_name", "") or ""

    # Folio
    if tipo_real == "diploma":
        folio = f"DIP-{req.id:06d}"
    else:
        folio = f"CON-{req.id:06d}"

    # Fechas numéricas dd/mm/aaaa
    def fmt_fecha_num(fecha):
        return fecha.strftime("%d/%m/%Y") if fecha else ""

    ctx = {
        # Identificación básica
        "folio": folio,
        "id": req.id,

        # Nombre / programa
        "nombre": nombre,          # 👈 para usar [nombre]
        "apellido": apellido,      # 👈 para usar [apellido]
        # Si algún día quieres el completo en el DOCX, se podría agregar otra key,
        # pero por ahora NO exponemos [nombre_completo] como pediste.
        "programa": programa,

        # Horas
        "horas": horas or "",

        # Fechas numéricas (para DC3 también)
        "vigencia_inicio": fmt_fecha_num(inicio),
        "vigencia_termino": fmt_fecha_num(fin),
        "fecha_inicio": fmt_fecha_num(inicio),
        "fecha_fin": fmt_fecha_num(fin),

        # Fechas en letra
        "vigencia_inicio_letra": _fecha_es_letra(inicio),
        "vigencia_termino_letra": _fecha_es_letra(fin),
        "fecha_inicio_letra": _fecha_es_letra(inicio),
        "fecha_fin_letra": _fecha_es_letra(fin),

        # Partes numéricas separadas (por si las quieres usar en DC3)
        "inicio_dia": inicio.day if inicio else "",
        "inicio_mes": inicio.month if inicio else "",
        "inicio_anio": inicio.year if inicio else "",

        "fin_dia": fin.day if fin else "",
        "fin_mes": fin.month if fin else "",
        "fin_anio": fin.year if fin else "",

        # Datos laborales / empresa (para DC3 / CPROEM)
        "curp": curp,
        "rfc": rfc,
        "puesto": puesto,
        "giro": giro,
        "razon_social": razon_social,

        # Verificación
        "token": token or "",
        "verify_url": verify_url or "",
    }

    return ctx

def _fecha_es_larga(fecha):
    if not fecha:
        return "—"

    # datetime -> date
    if isinstance(fecha, datetime):
        fecha = fecha.date()
    if not isinstance(fecha, date):
        try:
            fecha = localdate(fecha)
        except Exception:
            return str(fecha)

    dia = fecha.day
    mes = MESES_ES[fecha.month] if 1 <= fecha.month <= 12 else ""
    anio = fecha.year
    return f"{dia} de {mes} de {anio}"

# En views.py (donde ya tienes MESES_ES y _fecha_es_larga)

def _docx_placeholders_for_request(req, tipo_real: str, verify_url: str | None = None):
    """
    Placeholders disponibles para DOCX:

    Minúsculas:
      [fecha_inicio_letra]
      [fecha_fin_letra]

    MAYÚSCULAS:
      [fecha_inicio_letra_may]
      [fecha_fin_letra_may]
    """

    # Nombre y programa
    nombre = (getattr(req, "name", "") or "").strip()
    apellidos = (getattr(req, "lastname", "") or "").strip()
    nombre_completo = f"{nombre} {apellidos}".strip()

    program_obj = getattr(req, "program", None)
    programa = getattr(program_obj, "name", "") or ""

    # Fechas: preferimos Graduate.validity_*
    grad = getattr(req, "graduate", None)
    if grad and (getattr(grad, "validity_start", None) or getattr(grad, "validity_end", None)):
        inicio = getattr(grad, "validity_start", None)
        fin    = getattr(grad, "validity_end", None)
    else:
        inicio = getattr(req, "start_date", None)
        fin    = getattr(req, "end_date", None)

    # Horas
    horas = None
    if grad and hasattr(grad, "hours") and grad.hours is not None:
        horas = grad.hours
    elif hasattr(req, "hours") and req.hours is not None:
        horas = req.hours
    elif program_obj is not None and hasattr(program_obj, "hours") and program_obj.hours is not None:
        horas = program_obj.hours

    # Folio según tipo
    if tipo_real == "diploma":
        folio = f"DIP-{req.id:06d}"
    else:
        folio = f"CON-{req.id:06d}"

    # Fechas en partes
    def _parts(d):
        if not d:
            return ("", "", "")
        try:
            return (f"{d.day:02d}", f"{d.month:02d}", str(d.year))
        except Exception:
            return ("", "", "")

    inicio_dia, inicio_mes, inicio_anio = _parts(inicio)
    fin_dia, fin_mes, fin_anio = _parts(fin)

    # Formatos de fecha
    if inicio:
        inicio_letra = _fecha_es_larga(inicio) or ""
        inicio_letra_may = inicio_letra.upper()
    else:
        inicio_letra = ""
        inicio_letra_may = ""

    if fin:
        fin_letra = _fecha_es_larga(fin) or ""
        fin_letra_may = fin_letra.upper()
    else:
        fin_letra = ""
        fin_letra_may = ""

    # Diccionario principal (minúsculas)
    ctx = {
        "nombre": nombre,
        "apellido": apellidos,
        "nombre_completo": nombre_completo,

        "curp": getattr(req, "curp", "") or "",
        "rfc": getattr(req, "rfc", "") or "",
        "puesto": getattr(req, "job_title", "") or "",
        "giro": getattr(req, "industry", "") or "",
        "razon_social": getattr(req, "business_name", "") or "",
        "email": getattr(req, "email", "") or "",

        "programa": programa,
        "horas": str(horas) if horas is not None else "",

        "fecha_inicio": inicio.strftime("%d/%m/%Y") if inicio else "",
        "fecha_fin":    fin.strftime("%d/%m/%Y")     if fin else "",

        "fecha_inicio_letra": inicio_letra,
        "fecha_fin_letra":    fin_letra,

        # ¡NUEVOS PLACEHOLDERS!
        "fecha_inicio_letra_may": inicio_letra_may,
        "fecha_fin_letra_may":    fin_letra_may,

        "inicio_dia": inicio_dia,
        "inicio_mes": inicio_mes,
        "inicio_anio": inicio_anio,
        "fin_dia": fin_dia,
        "fin_mes": fin_mes,
        "fin_anio": fin_anio,

        "folio": folio,
        "verificacion_url": verify_url or "",
    }

    # Alias en MAYÚSCULAS
    upper_aliases = {
        "NOMBRE": "nombre",
        "APELLIDOS": "apellido",
        "NOMBRE_COMPLETO": "nombre_completo",
        "CURP": "curp",
        "RFC": "rfc",
        "PUESTO": "puesto",
        "GIRO": "giro",
        "RAZON_SOCIAL": "razon_social",
        "EMAIL": "email",
        "PROGRAMA": "programa",
        "HORAS": "horas",

        "FECHA_INICIO": "fecha_inicio",
        "FECHA_FIN": "fecha_fin",

        # Alias para minúsculas
        "FECHA_INICIO_LETRA": "fecha_inicio_letra",
        "FECHA_FIN_LETRA":    "fecha_fin_letra",

        # Alias MAYÚSCULAS nuevos
        "FECHA_INICIO_LETRA_MAY": "fecha_inicio_letra_may",
        "FECHA_FIN_LETRA_MAY":    "fecha_fin_letra_may",

        "FOLIO": "folio",
        "VERIFICACION_URL": "verificacion_url",
    }

    for upper_key, lower_key in upper_aliases.items():
        ctx[upper_key] = ctx.get(lower_key, "")

    return ctx


def _render_docx_template_for_request(req, tipo_real: str, verify_url: str | None = None):
    """
    Rellena la plantilla DOCX del programa asociada a esta Request
    (diploma / dc3 / cproem) usando los datos del egresado y devuelve
    los bytes del DOCX resultante.

    Si no hay plantilla configurada, devuelve None.
    """
    import io
    import re
    try:
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.oxml.text.paragraph import CT_P
        from docx.shared import Mm
    except Exception as e:
        raise RuntimeError(f"python-docx no está instalado o falló al importarse: {e}")

    try:
        import qrcode
    except Exception:
        qrcode = None

    # Localizamos el TemplateAsset correcto (según programa + tipo)
    tpl = _get_docx_template_for_request(req, tipo_real)
    if not tpl or not getattr(tpl, "file", None):
        # No hay plantilla DOCX configurada para este programa/tipo
        return None

    # Cargamos el DOCX
    doc = Document(tpl.file)

    # Mapeo de placeholders -> valores (usa datos de Request/Graduate/Programa)
    mapping = _docx_placeholders_for_request(req, tipo_real, verify_url=verify_url)

    # Generamos bytes del QR si hay verify_url y librería disponible
    qr_bytes = None
    if verify_url and qrcode is not None:
        try:
            qr = qrcode.QRCode(version=None, box_size=10, border=1)
            qr.add_data(verify_url)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            buff = io.BytesIO()
            img.save(buff, format="PNG")
            qr_bytes = buff.getvalue()
        except Exception:
            qr_bytes = None

    def _get_qr_size_from_paragraph(paragraph):
        """
        Intenta deducir el tamaño del cuadro de texto / shape que
        contiene este párrafo, a partir del XML (VML o DrawingML).

        Devuelve (width_mm, height_mm) o None si no se puede deducir.
        """
        el = getattr(paragraph, "_element", None) or getattr(paragraph, "_p", None)
        if el is None:
            return None

        # Subimos por los padres buscando v:shape o wp:inline/wp:anchor
        current = el
        width_mm = height_mm = None
        while current is not None:
            tag = current.tag

            # Caso 1: VML <v:shape style="... width:84pt;height:67.5pt; ...">
            if tag.endswith("}shape") and "urn:schemas-microsoft-com:vml" in tag:
                style = current.get("style") or ""
                m_w = re.search(r"width:([0-9.]+)pt", style)
                m_h = re.search(r"height:([0-9.]+)pt", style)
                if m_w and m_h:
                    try:
                        w_pt = float(m_w.group(1))
                        h_pt = float(m_h.group(1))
                        # pt -> mm (1 in = 72 pt; 1 in = 25.4 mm)
                        width_mm = w_pt * 25.4 / 72.0
                        height_mm = h_pt * 25.4 / 72.0
                        break
                    except Exception:
                        pass

            # Caso 2: DrawingML <wp:inline>/<wp:anchor> con <wp:extent cx= cy=>
            if tag.endswith("}inline") or tag.endswith("}anchor"):
                # Buscamos cualquier hijo <wp:extent> o <a:ext>
                extent = None
                for child in current.iter():
                    if child.tag.endswith("}extent") or child.tag.endswith("}ext"):
                        extent = child
                        break
                if extent is not None:
                    cx = extent.get("cx") or extent.get("cx", None)
                    cy = extent.get("cy") or extent.get("cy", None)
                    try:
                        if cx is not None and cy is not None:
                            cx_val = float(cx)
                            cy_val = float(cy)
                            # EMU -> mm (1 in = 914400 EMUs)
                            width_mm = cx_val / 914400.0 * 25.4
                            height_mm = cy_val / 914400.0 * 25.4
                            break
                    except Exception:
                        pass

            parent = getattr(current, "getparent", None)
            current = parent() if callable(parent) else None

        if width_mm and height_mm:
            # Para evitar recortes usamos un 90% del tamaño disponible
            side = min(width_mm, height_mm) * 0.9
            return (side, side)

        return None

    def apply_mapping_to_paragraphs(paragraphs):
        """
        Word suele partir el texto en varios runs, por eso trabajamos
        a nivel de párrafo. Juntamos el texto, reemplazamos, y luego
        lo escribimos de vuelta en el primer run.

        Además, si el párrafo es solo "[qr]" o "[QR]", insertamos el
        código QR como imagen, intentando usar el tamaño del cuadro.
        """
        for p in paragraphs:
            if not getattr(p, "runs", None):
                continue

            original = "".join(run.text or "" for run in p.runs)
            stripped = original.strip()

            # Caso especial: párrafo que SOLO contiene [qr]
            if stripped in ("[qr]", "[QR]"):
                # Borramos el texto
                for run in p.runs:
                    run.text = ""

                if qr_bytes:
                    # Calculamos tamaño según el cuadro de texto / shape
                    size = _get_qr_size_from_paragraph(p)
                    try:
                        run = p.runs[0] if p.runs else p.add_run()
                        img_stream = io.BytesIO(qr_bytes)
                        if size:
                            w_mm, h_mm = size
                            run.add_picture(img_stream, width=Mm(w_mm), height=Mm(h_mm))
                        else:
                            # Tamaño por defecto razonable
                            run.add_picture(img_stream, width=Mm(22), height=Mm(22))
                    except Exception:
                        # Si algo falla al insertar la imagen, simplemente lo dejamos vacío
                        pass

                # Pasamos al siguiente párrafo
                continue

            # Reemplazos normales de texto
            new_text = original
            for key, value in mapping.items():
                placeholder = f"[{key}]"
                val = "" if value is None else str(value)
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, val)

            if new_text == original:
                continue

            first = True
            for run in p.runs:
                if first:
                    run.text = new_text
                    first = False
                else:
                    run.text = ""

    def iter_all_paragraphs(document):
        """
        Itera TODOS los <w:p> del documento:
          - cuerpo principal
          - tablas
          - cuadros de texto / shapes (w:txbxContent, VML)
          - encabezados y pies (incluyendo sus text boxes)
        """
        # Body (incluye párrafos dentro de text boxes anclados al body)
        for el in document.element.body.iter():
            if isinstance(el, CT_P):
                yield Paragraph(el, document)

        # Headers y footers completos
        for section in document.sections:
            for container in (section.header, section.footer):
                try:
                    if container is None:
                        continue
                    for el in container._element.iter():
                        if isinstance(el, CT_P):
                            yield Paragraph(el, document)
                except Exception:
                    continue

    # Aplicamos reemplazos a todos los párrafos del documento
    try:
        all_paragraphs = list(iter_all_paragraphs(doc))
        apply_mapping_to_paragraphs(all_paragraphs)
    except Exception:
        # Si algo falla, no rompemos todo el render
        pass

    # Devolvemos bytes del DOCX generado
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


# ============================================================
# Selección automática de plantilla DOCX
# ============================================================

def _get_docx_template_for_request(req, tipo_real: str):
    """
    Devuelve la DocxTemplate correcta para este Request+tipo.

    Ahora usa los FKs de Program:
      - program.docx_tpl_diploma
      - program.docx_tpl_dc3
      - program.docx_tpl_cproem

    y ya NO intenta adivinar por nombre de archivo ni por FileField.

    Lógica:
      1) Si el Program tiene plantilla específica para ese tipo, se usa esa.
      2) Si no, se intenta una plantilla global activa (DocxTemplate.tipo = tipo, is_active = True).
      3) Si no hay ninguna, se devuelve None y se usará el fallback HTML.
    """
    program = getattr(req, "program", None)
    if not program:
        # Esto sí es un error de programación: el Request debería tener program
        raise Exception("El Request no tiene programa asociado.")

    # Normalizar el tipo que llega
    tipo = (tipo_real or "").strip().lower()

    # Si viene algo genérico como "constancia", decidimos según el Program.constancia_type
    if tipo in ("constancia", "const"):
        if getattr(program, "constancia_type", None) == ConstanciaType.DC3:
            tipo = "dc3"
        else:
            # Por defecto tratamos como CPROEM
            tipo = "cproem"

    # ------- DIPLOMA -------
    if tipo == "diploma":
        # 1) Plantilla específica del programa (si existe)
        tpl = getattr(program, "docx_tpl_diploma", None)
        if tpl:
            return tpl  # DocxTemplate

    # ------- DC3 -------
    if tipo == "dc3":
        # 1) Plantilla específica del programa (si existe)
        tpl = getattr(program, "docx_tpl_dc3", None)
        if tpl:
            return tpl  # DocxTemplate

    # ------- CPROEM -------
    if tipo == "cproem":
        # 1) Plantilla específica del programa (si existe)
        tpl = getattr(program, "docx_tpl_cproem", None)
        if tpl:
            return tpl  # DocxTemplate

    # 2) Fallback opcional: buscar una plantilla DOCX global activa de ese tipo
    tpl = DocxTemplate.objects.filter(tipo=tipo, is_active=True).first()
    if tpl:
        return tpl

    # 3) Si no hay nada (ni específica ni global), devolvemos None
    #    y el flujo usará la plantilla HTML (WeasyPrint) sin lanzar excepción.
    return None


def _build_cproem_context(*args):
    """
    Construye el contexto para plantillas CPROEM.

    Acepta dos formas de llamada por compatibilidad:
      - _build_cproem_context(request, graduate)
      - _build_cproem_context(graduate, request)

    Devuelve un dict con muchas claves (incluye `qr_url`) que usa
    tu template `pdf_constancia_cproem*.html`.
    """
    # ---- Normalizar argumentos ----
    request = None
    graduate = None

    if len(args) == 2:
        a0, a1 = args
        # detectar si alguno parece un HttpRequest (tiene build_absolute_uri)
        if hasattr(a0, "build_absolute_uri") and not hasattr(a1, "build_absolute_uri"):
            request, graduate = a0, a1
        elif hasattr(a1, "build_absolute_uri") and not hasattr(a0, "build_absolute_uri"):
            request, graduate = a1, a0
        else:
            # ninguno tiene build_absolute_uri (o ambos lo tienen) -> intentar heurística
            # asumimos (graduate, request) por compatibilidad con tu código previo
            graduate, request = a0, a1
    elif len(args) == 1:
        # si solo pasaron uno, asumimos que es graduate (no hay request disponible)
        graduate = args[0]
        request = None
    else:
        raise TypeError("_build_cproem_context espera (request, graduate) o (graduate, request)")

    # ---- Helpers / extracción de datos ----
    def _get_program_from_graduate(g):
        # intenta obtener program desde atributos comunes
        try:
            return getattr(getattr(g, "request", None), "program", None) or getattr(g, "program", None)
        except Exception:
            return None

    def _formatear_fecha_es(d):
        if not d:
            return ""
        try:
            return d.strftime("%d/%m/%Y")
        except Exception:
            return str(d)

    program = _get_program_from_graduate(graduate)

    # ==== Nombre del alumno ====
    full_name = (
        getattr(graduate, "full_name", None)
        or getattr(graduate, "nombre_completo", None)
        or f"{getattr(graduate, 'name', '')} {getattr(graduate, 'lastname', '')}".strip()
        or str(graduate)
    )

    # ==== Fechas ====
    start = (
        getattr(graduate, "validity_start", None)
        or getattr(graduate, "start_date", None)
        or getattr(graduate, "fecha_inicio", None)
        or getattr(graduate, "inicio", None)
    )
    end = (
        getattr(graduate, "validity_end", None)
        or getattr(graduate, "end_date", None)
        or getattr(graduate, "fecha_fin", None)
        or getattr(graduate, "fin", None)
    )

    fecha_inicio_txt = _formatear_fecha_es(start)
    fecha_fin_txt = _formatear_fecha_es(end)

    # ==== Horas / curp ====
    horas = (
        getattr(program, "hours", None)
        or getattr(graduate, "hours", None)
        or 120
    )
    curp = getattr(graduate, "curp", "") or ""

    # ==== verificacion_url (preparar con request si está) ====
    verificacion_url = None
    try:
        if request is not None:
            # intenta URL específica si existe la view
            try:
                verificacion_url = request.build_absolute_uri(
                    reverse("alumnos:verificar_documento", args=[getattr(graduate, "pk", getattr(graduate, "id", ""))])
                )
            except Exception:
                # fallback a seguimiento (como tenías antes)
                verificacion_url = request.build_absolute_uri(f"/alumnos/seguimiento/{getattr(graduate, 'pk', getattr(graduate, 'id', ''))}/")
    except Exception:
        verificacion_url = None

    # Si no hay request/URL, hacemos un valor relativo (útil para tests)
    if not verificacion_url:
        verificacion_url = f"/alumnos/seguimiento/{getattr(graduate, 'pk', getattr(graduate, 'id', ''))}/"

    # ==== Generar QR PNG en memoria y convertir a data URI ==== 
    try:
        qr = qrcode.QRCode(box_size=10, border=1)
        qr.add_data(verificacion_url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buffer = io.BytesIO()
        img.save(buffer, format="PNG")
        buffer.seek(0)
        qr_data_uri = "data:image/png;base64," + b64encode(buffer.getvalue()).decode("utf-8")
    except Exception:
        # Si falla la generación de QR, dejamos None (template puede mostrar fallback)
        qr_data_uri = None

    # ==== Contexto final (muchos alias para compatibilidad con templates) ====
    ctx = {
        "graduate": graduate,

        # nombres / alias
        "full_name": full_name,
        "nombre": full_name,
        "nombre_completo": full_name,
        "alumno_nombre": full_name,
        "student_name": full_name,

        # programa
        "program": program,
        "programa": program,
        "programa_nombre": getattr(program, "name", "") if program else "",
        "program_name": getattr(program, "name", "") if program else "",

        # horas, curp
        "horas": horas,
        "hours": horas,
        "curp": curp,
        "clave_curp": curp,

        # fechas crudas y formateadas (varios alias)
        "start_date": start,
        "end_date": end,
        "vigencia_inicio": start,
        "vigencia_termino": end,
        "fecha_inicio": fecha_inicio_txt,
        "fecha_fin": fecha_fin_txt,
        "inicio": fecha_inicio_txt,
        "fin": fecha_fin_txt,
        "fecha_inicio_txt": fecha_inicio_txt,
        "fecha_fin_txt": fecha_fin_txt,

        # verificación + QR (varios nombres para compatibilidad)
        "verificacion_url": verificacion_url,
        "verification_url": verificacion_url,
        "qr_data_uri": qr_data_uri,
        "qr_image": qr_data_uri,
        "qr": qr_data_uri,
        "qr_src": qr_data_uri,
        "qr_url": qr_data_uri,   # ← EXACTO: asegura que {{ qr_url }} en tu template funcione
    }

    return ctx


# Campos que permitimos actualizar desde el form/JS
ALLOWED_EGRESADO_FIELDS = [
    "nombre", "apellido_paterno", "apellido_materno",
    "curp", "rfc", "nss", "telefono", "email",
    "status", "notas",
]

# === FONDOS DE DIPLOMAS POR CÓDIGO (abreviado) ===
# El archivo PNG debe estar en: static/img/diplomas/
DIPLOMA_BG_BY_CODE = {
    "DAIE": "DAIEN.png",  # Diplomado en Administración de Instituciones Educativas
    "DAP":  "DAPN.png",
    "DCG":  "DCGN.png",
    "DDE":  "DDE.png",
    "DEIN": "DEIN.png",
    "DEM":  "DEMN.png",
    "DET":  "DETN.png",
    "DIC":  "DICN.png",
    "DLI":  "DLIN.png",
    "DNDI": "DNDI.png",
    "DNE":  "DNEN.png",
    "DNP":  "DNPN.png",
    "DNS":  "DNSN.png",
    "DOT":  "DOT.png",
    "DPI":  "DPIE.png",
    "DTC":  "DTCN.png",
    "DTS":  "DTSN.png",
    "DMVC": "DMVC.png"
}

def get_diploma_bg_for_code(code: str) -> str:
    """
    Devuelve la URL del fondo para el código de diplomado.

    1) Busca en DiplomaBackground (BD).
    2) Si no encuentra, regresa un PNG por defecto en static.
    """
    code = (code or "").strip().upper()
    if not code:
        return static("img/diplomas/DEFAULT.png")  # pon aquí tu fondo genérico

    try:
        bg = DiplomaBackground.objects.get(code=code, is_active=True)
        # si usas WeasyPrint con rutas absolutas, quizá quieras bg.image.path
        return bg.image.path   # o bg.image.url según como lo uses
    except DiplomaBackground.DoesNotExist:
        return static("img/diplomas/DEFAULT.png")
# ====================== Endpoint well-known (Chrome DevTools) ======================
@require_GET
def wellknown_devtools(request):
    """
    Atiende /.well-known/appspecific/com.chrome.devtools.json
    Devolver un JSON mínimo es suficiente para que no falle el import en urls.py.
    """
    return JsonResponse({
        "name": "CES System DevTools",
        "status": "ok"
    })

# ====================== Helpers de autorización ======================

def _in_admin_group(user) -> bool:
    try:
        return user.groups.filter(name="AdministradoresCES").exists()
    except Exception:
        return False

def _admin_allowed(user) -> bool:
    """
    Versión relajada:
    Por ahora, cualquier usuario autenticado y activo puede usar el panel
    (no diferenciamos entre superuser / staff / grupos).
    """
    return getattr(user, "is_authenticated", False) and getattr(user, "is_active", False)


class StaffRequiredMixin(UserPassesTestMixin):
    def test_func(self):
        return _admin_allowed(self.request.user)

    def handle_no_permission(self):
        return redirect("administracion:login")


class OTPRequiredMixin(StaffRequiredMixin):
    """Exige 2FA (TOTP) para acceder al panel."""
    @method_decorator(otp_required(login_url=reverse_lazy("two_factor:login")))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)


# =========================== Home ===========================

class AdminHomeView(LoginRequiredMixin, TemplateView):
    template_name = "administracion/home.html"
    login_url = "administracion:login"

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx["total_requests"] = Request.objects.count()
        ctx["pending"] = Request.objects.filter(status="pending").count()
        ctx["accepted"] = Request.objects.filter(status="accepted").count()
        ctx["rejected"] = Request.objects.filter(status="rejected").count()
        ctx["total_graduates"] = Graduate.objects.count()

        q = (self.request.GET.get("q") or "").strip()
        ctx["q"] = q
        qs = Request.objects.select_related("program").order_by("-sent_at")
        if q:
            qs = qs.filter(
                Q(email__icontains=q) |
                Q(curp__icontains=q) |
                Q(name__icontains=q) |
                Q(lastname__icontains=q) |
                Q(program__name__icontains=q)
            )

        ctx["recent_requests"] = qs[:20]
        ctx["recent_logs"] = AdminAccessLog.objects.select_related("user").order_by("-created_at")[:12]
        ctx["now"] = timezone.now()

        try:
            ctx["user_groups"] = list(self.request.user.groups.values_list("name", flat=True))
        except Exception:
            ctx["user_groups"] = None

        ctx["csrf_token"] = get_token(self.request)
        return ctx


# ======================= Login / Logout =======================

class AdminLoginView(FormView):
    template_name = "administracion/login.html"
    form_class = AdminLoginForm
    success_url = reverse_lazy("administracion:home")

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        email = self.request.POST.get("email") or ""
        fails = failed_attempts_count(self.request, username_or_email=email)
        ctx["show_hcaptcha"] = fails >= getattr(settings, "HCAPTCHA_THRESHOLD_FAILS", 3)
        ctx["HCAPTCHA_SITE_KEY"] = getattr(settings, "HCAPTCHA_SITE_KEY", "")
        return ctx

    def dispatch(self, request, *args, **kwargs):
        if request.user.is_authenticated:
            return redirect("administracion:home")
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        email = form.cleaned_data["email"]
        password = form.cleaned_data["password"]
        remember = form.cleaned_data["remember_me"]
        needs_captcha = failed_attempts_count(
            self.request, username_or_email=email
        ) >= getattr(settings, "HCAPTCHA_THRESHOLD_FAILS", 3)

        if needs_captcha:
            token = self.request.POST.get("h-captcha-response", "")
            if not token:
                form.add_error(None, "Por favor, completa la verificación.")
                return self.form_invalid(form)
            secret = getattr(settings, "HCAPTCHA_SECRET_KEY", "")
            try:
                resp = requests.post(
                    "https://hcaptcha.com/siteverify",
                    data={
                        "secret": secret,
                        "response": token,
                        "remoteip": self.request.META.get("REMOTE_ADDR", ""),
                    },
                    timeout=5,
                )
                data = resp.json()
                if not data.get("success"):
                    form.add_error(None, "Verificación fallida. Intenta de nuevo.")
                    return self.form_invalid(form)
            except Exception:
                form.add_error(None, "No pudimos validar el captcha. Intenta de nuevo.")
                return self.form_invalid(form)

        try:
            user = User.objects.get(email__iexact=email)
        except User.DoesNotExist:
            user = None

        auth_user = authenticate(
            self.request,
            username=getattr(user, "username", email),
            password=password
        ) if user else None

        if not auth_user or not auth_user.is_active or not _admin_allowed(auth_user):
            form.add_error(None, "Credenciales inválidas.")
            return self.form_invalid(form)

        login(self.request, auth_user)

        if remember:
            self.request.session.set_expiry(int(timedelta(hours=8).total_seconds()))
            self.request.session["remember_me"] = True
        else:
            self.request.session.set_expiry(0)
            self.request.session["remember_me"] = False

        return super().form_valid(form)

class AdminLogoutView(TemplateView):
    def get(self, request, *args, **kwargs):
        logout(request)
        return redirect("administracion:login")


# ================ Password reset (plantillas personalizadas) ================

class AdminPasswordResetView(PasswordResetView):
    template_name = "administracion/password_reset.html"
    email_template_name = "administracion/password_reset_email.txt"
    subject_template_name = "administracion/password_reset_subject.txt"
    success_url = reverse_lazy("administracion:password_reset_done")

class AdminPasswordResetDoneView(PasswordResetDoneView):
    template_name = "administracion/password_reset_done.html"

class AdminPasswordResetConfirmView(PasswordResetConfirmView):
    template_name = "administracion/password_reset_confirm.html"
    success_url = reverse_lazy("administracion:password_reset_complete")

class AdminPasswordResetCompleteView(PasswordResetCompleteView):
    template_name = "administracion/password_reset_complete.html"


# =========================== Endpoints AJAX (cuenta) ===========================
@login_required
def home(request):
    q = (request.GET.get("q") or "").strip()
    order = (request.GET.get("order") or "date_desc").strip()

    # Base queryset de solicitudes
    req_qs = Request.objects.select_related("program")

    if q:
        req_qs = req_qs.filter(
            Q(email__icontains=q) |
            Q(name__icontains=q) |
            Q(lastname__icontains=q) |
            Q(curp__icontains=q) |
            Q(program__name__icontains=q) |
            Q(program__abbreviation__icontains=q)
        )

    # ---------- ORDEN ----------
    if order == "date_asc":
        req_qs = req_qs.order_by("sent_at","id")
    elif order == "name_asc":
        req_qs = req_qs.order_by("name", "lastname","id")
    elif order == "name_desc":
        req_qs = req_qs.order_by("-name", "-lastname","-id")
    else:
        # default: más recientes primero
        order = "date_desc"
        req_qs = req_qs.order_by("-sent_at","-id")

    recent_requests = req_qs[:50]

    # KPIs
    total_requests = Request.objects.count()
    pending = Request.objects.filter(status="pending").count()
    accepted = Request.objects.filter(status="accepted").count()
    rejected = Request.objects.filter(status="rejected").count()

    # Logs de seguridad
    recent_logs = AdminAccessLog.objects.select_related("user")[:10]

    ctx = {
        "q": q,
        "order": order,
        "recent_requests": recent_requests,
        "total_requests": total_requests,
        "pending": pending,
        "accepted": accepted,
        "rejected": rejected,
        "recent_logs": recent_logs,
        "now": timezone.now(),
    }
    return render(request, "administracion/home.html", ctx)

@require_POST
def password_change_inline(request):
    if not request.user.is_authenticated:
        return JsonResponse({"ok": False, "error": "No autenticado."}, status=401)

    p1 = (request.POST.get("new_password1") or "").strip()
    p2 = (request.POST.get("new_password2") or "").strip()

    if not p1 or not p2:
        return JsonResponse({"ok": False, "error": "Completa ambos campos."}, status=400)
    if p1 != p2:
        return JsonResponse({"ok": False, "error": "Las contraseñas no coinciden."}, status=400)
    if len(p1) < 8:
        return JsonResponse({"ok": False, "error": "La contraseña debe tener al menos 8 caracteres."}, status=400)

    user = request.user
    user.set_password(p1)
    user.save()
    update_session_auth_hash(request, user)
    return JsonResponse({"ok": True})

@login_required
@csrf_protect
@require_POST
def create_admin_inline(request):
    """
    Crea un nuevo usuario admin desde el modal inline.
    Siempre responde JSON, incluso si hay errores internos.
    """
    try:
        me = request.user
        if not _admin_allowed(me):
            # Respuesta amigable para el modal (no 403 duro)
            return JsonResponse({"ok": False, "error": "No autorizado."}, status=200)

        email      = (request.POST.get("email") or "").strip().lower()
        username   = (request.POST.get("username") or "").strip()
        first_name = (request.POST.get("first_name") or "").strip()
        last_name  = (request.POST.get("last_name") or "").strip()
        pw1        = (request.POST.get("password1") or "").strip()
        pw2        = (request.POST.get("password2") or "").strip()

        # --------- Validaciones de negocio ---------
        if not email or "@" not in email:
            return JsonResponse({"ok": False, "error": "Email inválido."}, status=200)

        if pw1 != pw2:
            return JsonResponse({"ok": False, "error": "Las contraseñas no coinciden."}, status=200)

        if len(pw1) < 8:
            return JsonResponse(
                {"ok": False, "error": "La contraseña debe tener al menos 8 caracteres."},
                status=200,
            )

        # Generar username si viene vacío
        if not username:
            base = slugify(email.split("@")[0]) or "admin"
            candidate, i = base, 1
            while User.objects.filter(username=candidate).exists():
                i += 1
                candidate = f"{base}{i}"
            username = candidate

        # Duplicados
        if User.objects.filter(email__iexact=email).exists():
            return JsonResponse(
                {"ok": False, "error": "Ya existe un usuario con ese email."},
                status=200,
            )

        if User.objects.filter(username=username).exists():
            return JsonResponse(
                {"ok": False, "error": "El nombre de usuario ya está en uso."},
                status=200,
            )

        # --------- Crear usuario (OJO: AdminUser no acepta first_name/last_name como kwargs) ---------
        u = User.objects.create_user(
            username=username,
            email=email,
            password=pw1,
            is_active=True,
        )

        # Si tu modelo tiene campos de nombre, los llenamos de forma segura
        if first_name or last_name:
            changed = False

            # Caso 1: tu modelo define first_name / last_name
            if hasattr(u, "first_name"):
                u.first_name = first_name
                changed = True
            if hasattr(u, "last_name"):
                u.last_name = last_name
                changed = True

            # Caso 2: modelo con un solo campo de nombre (name / full_name / nombre)
            if not changed:
                full = f"{first_name} {last_name}".strip()
                for attr in ("name", "full_name", "nombre"):
                    if hasattr(u, attr):
                        setattr(u, attr, full)
                        changed = True
                        break

            if changed:
                u.save()

        # Convertirlo en staff/admin
        if hasattr(u, "is_staff"):
            u.is_staff = True
            u.save(update_fields=["is_staff"])

        grp, _ = Group.objects.get_or_create(name="AdministradoresCES")
        u.groups.add(grp)

        return JsonResponse(
            {
                "ok": True,
                "user": {
                    "id": u.id,
                    "username": u.username,
                    "email": u.email,
                },
            },
            status=200,
        )

    except Exception as e:
        # Cualquier cosa rara (constraints, campos obligatorios extra, etc.)
        log.exception("Error en create_admin_inline")
        return JsonResponse(
            {
                "ok": False,
                "error": f"Error interno: {e.__class__.__name__}: {e}",
            },
            status=200,
        )

# ======================== Plantillas (portada) ========================

@login_required
def plantillas(request, req_id=None):
    generating = (
        Request.objects
        .filter(status='generating')
        .select_related('program')
        .order_by('-sent_at')
    )

    active_req = None
    if req_id:
        active_req = get_object_or_404(
            Request.objects.select_related('program'),
            pk=req_id
        )

    full_name = ""
    program_name = ""
    active_rfc = ""
    active_curp = ""
    active_job_title = ""
    active_industry = ""

    if active_req:
        first = getattr(active_req, "name", "") or ""
        last = getattr(active_req, "lastname", "") or ""
        full_name = f"{first} {last}".strip()
        prog = getattr(active_req, "program", None)
        program_name = getattr(prog, "name", "") if prog else ""
        active_rfc = getattr(active_req, "rfc", "") or ""
        active_curp = getattr(active_req, "curp", "") or ""
        active_job_title = getattr(active_req, "job_title", "") or ""
        active_industry = getattr(active_req, "industry", "") or ""

    ctx = {
        'generating_reqs': generating,
        'active_req': active_req,
        'full_name': full_name,
        'program_name': program_name,
        'active_rfc': active_rfc,
        'active_curp': active_curp,
        'active_job_title': active_job_title,
        'active_industry': active_industry,
        'diploma': {},
        'dc3': {},
        'cproem': {},
    }
    return render(request, 'administracion/plantillas.html', ctx)


# ============================= Egresados =============================

def _random_hex_25() -> str:
    return secrets.token_hex(13)[:25]

def _qr_png_data_url(payload: str, box_size: int = 12, border: int = 2) -> str:
    qr = qrcode.QRCode(version=None, box_size=box_size, border=border)
    qr.add_data(payload)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buff = io.BytesIO()
    img.save(buff, format="PNG")
    return "data:image/png;base64," + b64encode(buff.getvalue()).decode("ascii")

def _constancia_tipo(program) -> str:
    name = (getattr(program, "name", "") or "").lower()
    if "cproem" in name:
        return "cproem"
    return "dc3"

def _pick(*vals):
    for v in vals:
        if v:
            return v
    return None

def _get_tipo_from_request(request, default="diploma"):
    return (request.GET.get("tipo")
            or request.GET.get("type")
            or request.POST.get("tipo")
            or request.POST.get("type")
            or default)

def _norm_txt(s: str) -> str:
    s = (s or "").strip().lower()
    s = unicodedata.normalize("NFD", s)
    return "".join(ch for ch in s if unicodedata.category(ch) != "Mn")

def _map_admin_constancia_type(val) -> str:
    """
    Mapea un valor de constancia_type (enum/int/str) a 'cproem' o 'dc3'.
    """
    if val is None:
        return ""
    try:
        from .models import ConstanciaType as _CT
        raw = getattr(val, "value", val)
        if isinstance(raw, int):
            if getattr(_CT, "CEPROEM", None) == raw:
                return "cproem"
            if getattr(_CT, "DC3", None) == raw:
                return "dc3"
        sval = _norm_txt(str(raw))
        if "cproem" in sval or "ceproem" in sval:
            return "cproem"
        if "dc3" in sval:
            return "dc3"
    except Exception:
        sval = _norm_txt(str(getattr(val, "value", val)))
        if "cproem" in sval or "ceproem" in sval:
            return "cproem"
        if "dc3" in sval:
            return "dc3"
    return ""
def _constancia_kind_by_fields(req) -> str | None:
    """
    Aplica tu regla por presencia de campos.
    - Solo CURP (y sin RFC/puesto/giro/razón_social) => 'cproem'
    - Si hay cualquiera de los otros => 'dc3'
    Devuelve None si no puede decidir.
    """
    if not req:
        return None
    curp = (getattr(req, "curp", "") or "").strip()
    rfc = (getattr(req, "rfc", "") or "").strip()
    job = (getattr(req, "job_title", "") or "").strip()
    ind = (getattr(req, "industry", "") or "").strip()
    biz = (getattr(req, "business_name", "") or "").strip()

    has_extras = any([rfc, job, ind, biz])
    if curp and not has_extras:
        return "cproem"
    if has_extras:
        return "dc3"
    return None

def constancia_kind_for_request(obj) -> str:
    """
    Decide 'cproem' o 'dc3'.
    - Si obj es Request, primero respetamos Program si lo define claramente.
    - Regla de negocio: si SOLO hay CURP (y NO hay rfc, job_title, industry, business_name) => cproem.
      Si existen cualquiera de esos campos => dc3.
    """
    try:
        program = getattr(obj, "program", None)
    except Exception:
        program = None

    # 1) Lo que diga el programa (si aplica)
    kind = _program_constancia_kind(program)  # tu helper existente: 'dc3' | 'cproem'

    # 2) Validación por datos del alumno (regla tuya)
    curp = getattr(obj, "curp", None) if hasattr(obj, "curp") else None
    rfc = getattr(obj, "rfc", None) if hasattr(obj, "rfc") else None
    job = getattr(obj, "job_title", None) if hasattr(obj, "job_title") else None
    giro = getattr(obj, "industry", None) if hasattr(obj, "industry") else None
    biz  = getattr(obj, "business_name", None) if hasattr(obj, "business_name") else None

    # Si solo hay CURP y no hay nada más laboral → CPROEM (gana la regla de negocio)
    if curp and not any([rfc, job, giro, biz]):
        return "cproem"

    return kind or "dc3"

def _program_constancia_kind(program) -> str:
    """
    Devuelve 'cproem' o 'dc3' según:
    1) constancia_type en el propio objeto (si existe)
    2) cross-lookup en administracion.Program por code o name
    3) heurística por texto (name/code/abbreviation)
    4) fallback 'dc3'
    """
    if not program:
        return "dc3"

    # 1) Si el objeto trae constancia_type directo (caso ProgramAdmin)
    ct = getattr(program, "constancia_type", None)
    kind = _map_admin_constancia_type(ct)
    if kind:
        return kind

    # 2) Cross-lookup contra ProgramAdmin si lo tenemos disponible
    try:
        from .models import Program as ProgramAdmin
        code = getattr(program, "code", None)
        name = getattr(program, "name", None)

        cand = None
        if code:
            cand = ProgramAdmin.objects.filter(code__iexact=code).first()
        if not cand and name:
            cand = ProgramAdmin.objects.filter(name__iexact=name).first()

        if cand:
            kind = _map_admin_constancia_type(getattr(cand, "constancia_type", None))
            if kind:
                return kind
    except Exception:
        pass  # si falla el lookup, seguimos a heurística

    # 3) Heurística por texto (soporta CPROEM/CEPROEM)
    for attr in ("name", "code", "abbreviation"):
        sval = _norm_txt(getattr(program, attr, ""))
        if "cproem" in sval or "ceproem" in sval:
            return "cproem"

    # 4) Fallback
    return "dc3"

# Alias usado por el resto del código
def _constancia_tipo(program) -> str:
    return _program_constancia_kind(program)

@login_required
def egresados(request, req_id=None):
    """
    Panel de egresados.

    - EN PROCESO: solicitudes con status "generating".
    - FINALIZADOS: solicitudes que YA tienen Graduate asociado, sin depender del status.
    """

    # -----------------------
    #  Solicitudes EN PROCESO
    # -----------------------
    generating = (
        Request.objects
        .filter(status="generating")
        .select_related("program")
        .order_by("-sent_at")
    )

    # -----------------------
    #  Solicitudes FINALIZADAS
    #   (todas las que ya tienen Graduate, sin importar status)
    # -----------------------
    finished = (
        Request.objects
        .filter(graduate__isnull=False)
        .select_related("program", "graduate")
        .order_by("-graduate__completion_date", "-sent_at", "-id")
    )

    # Alumno activo (si viene un id en la URL)
    active_req = (
        Request.objects
        .select_related("program", "graduate")
        .filter(pk=req_id)
        .first()
        if req_id
        else None
    )

    # Datos básicos para el panel derecho
    program_name = None
    req_start = None
    req_end = None
    curp = rfc = job_title = industry = business_name = None
    hours = None  # 🔹 NUEVO: intensidad por defecto

    if active_req:
        program = getattr(active_req, "program", None)
        program_name = getattr(program, "name", None)

        grad = getattr(active_req, "graduate", None)

        # Preferimos las fechas que estén en Graduate
        if grad and (getattr(grad, "validity_start", None) or getattr(grad, "validity_end", None)):
            req_start = getattr(grad, "validity_start", None)
            req_end = getattr(grad, "validity_end", None)
        else:
            # Compatibilidad con la Request original
            req_start = getattr(active_req, "start_date", None)
            req_end = getattr(active_req, "end_date", None)

        # 🔹 NUEVO: intensidad desde Graduate (si existe)
        if grad:
            hours = getattr(grad, "hours", None)

        curp = getattr(active_req, "curp", None)
        rfc = getattr(active_req, "rfc", None)
        job_title = getattr(active_req, "job_title", None)
        industry = getattr(active_req, "industry", None)
        business_name = getattr(active_req, "business_name", None)

    # Tipo de constancia (dc3 / cproem) para el alumno activo
    const_kind = constancia_kind_for_request(active_req) if active_req else "dc3"
    is_dc3 = const_kind == "dc3"
    is_cproem = const_kind == "cproem"
    const_label = "DC-3" if is_dc3 else "CPROEM"

    ctx = {
        # Listas de la columna izquierda
        "generating_reqs": generating,   # en proceso
        "finished_reqs": finished,       # finalizados (por Graduate)

        # Alumno seleccionado
        "active_req": active_req,
        "program_name": program_name,
        "req_start": req_start,
        "req_end": req_end,
        "curp": curp,
        "rfc": rfc,
        "job_title": job_title,
        "industry": industry,
        "business_name": business_name,
        "hours": hours,  # 🔹 NUEVO: para usar en egresados.html

        # Info de tipo de constancia
        "is_dc3": is_dc3,
        "is_cproem": is_cproem,
        "constancia_label": const_label,
        "const_kind": const_kind,

        # Dummies para plantillas
        "diploma": {},
        "constancia": {},

        # Para el JS
        "csrf_token": get_token(request),
    }
    return render(request, "administracion/egresados.html", ctx)


@csp_update({
    'img-src': ("'self'", "data:", "blob:"),
    'style-src': ("'self'", "https://fonts.googleapis.com", csp.NONCE),
})
@require_GET
def egresados_preview_pdf(request, req_id: int):
    """
    Genera un PDF de vista previa (seguro) del diploma/constancia de la Request.

    - ?tipo=diploma|constancia
    - Usa primero la plantilla DOCX con variables (si existe).
    - Si no hay DOCX o falla, hace fallback a la plantilla HTML (WeasyPrint).
    """
    try:
        from weasyprint import HTML
    except Exception:
        HTML = None

    req = get_object_or_404(Request.objects.select_related("program"), pk=req_id)

    tipo = (request.GET.get("tipo") or "diploma").lower()
    if tipo not in ("diploma", "constancia"):
        return JsonResponse({"ok": False, "error": "Tipo inválido."}, status=400)

    # diploma | dc3 | cproem
    tipo_real = "diploma" if tipo == "diploma" else constancia_kind_for_request(req)

    # --- 1) Intentar DOCX con variables -> PDF ---
    pdf_bytes = None
    docx_bytes = None
    try:
        # Publicamos (o reutilizamos) token y armamos URL de verificación
        doc_token = _ensure_published_token(req, tipo_real)
        verify_url = request.build_absolute_uri(
            reverse("administracion:verificar_token", args=[doc_token.token])
        )

        docx_bytes = _render_docx_template_for_request(
            req,
            tipo_real,
            verify_url=verify_url,
        )
    except Exception as e:
        log.exception(
            "Error usando plantilla DOCX en preview %s #%s: %s",
            tipo_real,
            req.id,
            e,
        )
        docx_bytes = None

    if docx_bytes:
        try:
            class _U:
                def __init__(self, b):
                    self._b = b

                def chunks(self, csize=64 * 1024):
                    mv = memoryview(self._b)
                    for i in range(0, len(mv), csize):
                        yield mv[i : i + csize]

            raw_pdf = _office_to_pdf_bytes(
                _U(docx_bytes),
                name_hint=f"{tipo_real}-preview-{req.id}.docx",
            )

            # Misma "super seguridad": raster + permisos + marca de agua
            pdf_bytes = _pdf_secure(
                raw_pdf,
                user_pwd="",
                watermark_text="CES · Vista previa",
            )
        except Exception as e:
            log.exception(
                "Error DOCX→PDF en preview %s #%s: %s",
                tipo_real,
                req.id,
                e,
            )
            pdf_bytes = None

    # --- 2) Fallback: HTML + WeasyPrint ---
    if pdf_bytes is None:
        if HTML is None:
            return JsonResponse(
                {"ok": False, "error": "WeasyPrint no está instalado."},
                status=500,
            )

        try:
            tpl, ctx = _render_ctx_and_template(
                request,
                req,
                tipo_real,
                use_published_if_exists=False,
            )
            html = render_to_string(tpl, ctx, request=request)
            base_url = request.build_absolute_uri("/")
            raw_pdf = HTML(string=html, base_url=base_url).write_pdf()

            pdf_bytes = _pdf_secure(
                raw_pdf,
                user_pwd="",
                watermark_text="CES · Vista previa",
            )
        except Exception as e:
            log.exception(
                "Fallo generando vista previa HTML para %s #%s: %s",
                tipo_real,
                req.id,
                e,
            )
            return JsonResponse(
                {"ok": False, "error": f"Error generando vista previa: {e}"},
                status=500,
            )

    # Mostramos el PDF en el navegador (no como descarga forzada)
    return HttpResponse(pdf_bytes, content_type="application/pdf")


@login_required
def egresado_update(request, req_id):
    """
    Actualiza, vía AJAX, los datos reales del egresado.
    - La URL recibe el id de alumnos.Request (req_id).
    - Los datos “fuertes” se guardan en administracion.Graduate.
    - En lo posible, también se actualiza alumnos.Request si tiene esos campos.
    """
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])

    # 1) Localizar la solicitud original
    req = get_object_or_404(Request, pk=req_id)

    # 2) Localizar (o crear) el Graduate vinculado
    grad, _created = Graduate.objects.get_or_create(
        request=req,
        defaults={
            "name": getattr(req, "name", "") or "",
            "lastname": getattr(req, "lastname", "") or "",
            "email": getattr(req, "email", "") or "",
            "curp": getattr(req, "curp", "") or "",
        },
    )

    changed = {}
    grad_dirty = False
    req_dirty = False

    # --------- Campos de texto (nombre, curp, etc.) ---------
    # Cada entrada: nombre_en_POST -> lista de posibles destinos (objeto, campo)
    text_fields = {
        "name":         [("grad", "name"),         ("req", "name")],
        "lastname":     [("grad", "lastname"),     ("req", "lastname")],
        "curp":         [("grad", "curp"),         ("req", "curp")],
        "job_title":    [("grad", "job_title"),    ("req", "job_title")],
        "industry":     [("grad", "industry"),     ("req", "industry")],
        "business_name":[("grad", "business_name"),("req", "business_name")],
        # Si tu modelo Request tiene rfc, lo actualizamos ahí:
        "rfc":          [("req", "rfc")],
    }

    for post_name, targets in text_fields.items():
        raw_val = (request.POST.get(post_name) or "").strip()
        db_val = raw_val or None  # guardamos None si viene vacío

        for obj_name, field in targets:
            obj = grad if obj_name == "grad" else req
            if not hasattr(obj, field):
                continue

            old_val = getattr(obj, field) or ""
            if old_val != raw_val:
                setattr(obj, field, db_val)
                # solo registramos una vez el cambio por campo
                if post_name not in changed:
                    changed[post_name] = f"{old_val or '—'} → {raw_val or '—'}"

                if obj_name == "grad":
                    grad_dirty = True
                else:
                    req_dirty = True
                # pasamos al siguiente campo POST
                break

    # --------- Fechas (inicio / fin) ---------
    # POST: start_date / end_date
    # Graduate: validity_start / validity_end
    # Request: (si existen) start_date / end_date
    date_fields = {
        "start_date": [("grad", "validity_start"), ("req", "start_date")],
        "end_date":   [("grad", "validity_end"),   ("req", "end_date")],
    }

    for post_name, targets in date_fields.items():
        raw = (request.POST.get(post_name) or "").strip()

        if not raw:
            new_date = None
        else:
            new_date = parse_date(raw)  # espera YYYY-MM-DD
            if new_date is None:
                return HttpResponseBadRequest(f"Fecha inválida para {post_name}: {raw}")

        for obj_name, field in targets:
            obj = grad if obj_name == "grad" else req
            if not hasattr(obj, field):
                continue

            old_date = getattr(obj, field)
            if old_date != new_date:
                setattr(obj, field, new_date)

                if post_name not in changed:
                    changed[post_name] = f"{old_date or '—'} → {new_date or '—'}"

                if obj_name == "grad":
                    grad_dirty = True
                else:
                    req_dirty = True
                break

    # --------- HOURS (intensidad) ---------
    hours_raw = (request.POST.get("hours") or "").strip()
    if hours_raw:
        try:
            new_hours = int(hours_raw)
            if new_hours <= 0:
                new_hours = None
        except ValueError:
            return HttpResponseBadRequest(f"Valor inválido para horas: {hours_raw}")
    else:
        new_hours = None

    old_hours = getattr(grad, "hours", None)
    if old_hours != new_hours:
        grad.hours = new_hours
        grad_dirty = True
        changed["hours"] = f"{old_hours or '—'} → {new_hours or '—'}"

    # --------- Guardar si hubo cambios ---------
    if grad_dirty:
        grad.save()
    if req_dirty:
        req.save()

    return JsonResponse({
        "ok": True,
        "changed": changed,
    })


def egresados_req(request, req_id):
    active_req = get_object_or_404(Request, pk=req_id)

    # 👉 traemos (si existe) el Graduate ligado a la solicitud
    graduate = Graduate.objects.filter(request=active_req).first()

    # ----------------- DATOS BÁSICOS -----------------
    # si en Graduate hay dato, se usa; si no, se cae a Request
    curp = None
    if graduate and graduate.curp:
        curp = graduate.curp
    else:
        curp = active_req.curp

    job_title = None
    if graduate and graduate.job_title:
        job_title = graduate.job_title
    else:
        job_title = getattr(active_req, "job_title", None)

    industry = None
    if graduate and graduate.industry:
        industry = graduate.industry
    else:
        industry = getattr(active_req, "industry", None)

    business_name = None
    if graduate and graduate.business_name:
        business_name = graduate.business_name
    else:
        business_name = getattr(active_req, "business_name", None)

    # ----------------- FECHAS -----------------
    # 👈 AQUÍ ES LO IMPORTANTE
    if graduate and graduate.validity_start:
        req_start = graduate.validity_start
    else:
        req_start = getattr(active_req, "start_date", None)

    if graduate and graduate.validity_end:
        req_end = graduate.validity_end
    else:
        req_end = getattr(active_req, "end_date", None)

    # ... resto de tu lógica: generar diploma_data, constancia_data, etc. ...

    context = {
        "active_req": active_req,
        "curp": curp,
        "job_title": job_title,
        "industry": industry,
        "business_name": business_name,
        "req_start": req_start,
        "req_end": req_end,
        # + todo lo demás que ya mandabas:
        # "diploma": diploma_data,
        # "constancia": constancia_data,
        # "is_dc3": is_dc3,
        # "constancia_label": constancia_label,
        # "const_kind": const_kind,
        # "generating_reqs": generating_reqs,
        # etc...
    }

    return render(request, "administracion/egresados.html", context)

@login_required
def egresado_update_inline(request, req_id: int):
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])

    incoming = {k: v for k, v in request.POST.items() if k in ALLOWED_EGRESADO_FIELDS}

    if Egresado is not None:
        obj = get_object_or_404(Egresado, pk=req_id)
        for field, value in incoming.items():
            if hasattr(obj, field):
                setattr(obj, field, value)
        obj.save()
        return JsonResponse({"ok": True, "id": obj.pk, "saved": incoming})

    return JsonResponse({
        "ok": True,
        "id": req_id,
        "saved": incoming,
        "note": "Egresado model no disponible; stub.",
    })
    
def doc_send(request, req_id):
    """
    Genera el documento (preferiblemente desde DOCX con variables),
    publica el token de verificación y envía el PDF vía correo al alumno.
    Actualiza tracking (emailed / finalizado).
    """
    import logging
    from django.core.mail import EmailMultiAlternatives
    from django.utils import timezone
    from django.http import JsonResponse
    from django.shortcuts import get_object_or_404
    from django.templatetags.static import static as static_url
    from django.conf import settings
    from email.mime.image import MIMEImage
    from django.contrib.staticfiles import finders
    from django.utils.html import escape
    from django.template.loader import render_to_string
    from django.urls import reverse

    log = logging.getLogger(__name__)

    req = get_object_or_404(
        Request.objects.select_related("program", "graduate"),
        pk=req_id,
    )

    # ----------------------------------
    # Determinar tipo real (diploma/cproem/dc3)
    # ----------------------------------
    tipo_query = _get_tipo_from_request(request, default="diploma")
    if tipo_query not in ("diploma", "constancia"):
        return JsonResponse({"ok": False, "error": "Tipo inválido."}, status=400)

    tipo_real = "diploma" if tipo_query == "diploma" else constancia_kind_for_request(req)
    log.info("[doc_send] tipo_query=%s → tipo_real=%s (req=%s)", tipo_query, tipo_real, req.id)

    # ----------------------------------
    # Validar email del alumno
    # ----------------------------------
    if not req.email or not req.email.strip():
        return JsonResponse(
            {"ok": False, "error": "La solicitud no tiene correo electrónico capturado."},
            status=400,
        )

    # ----------------------------------
    # Publicar token / verificar URL
    # ----------------------------------
    doc = _ensure_published_token(req, tipo_real)
    verify_url = request.build_absolute_uri(
        reverse("administracion:verificar_token", args=[doc.token])
    )

    # ----------------------------------
    # Generar PDF desde DOCX (preferencia)
    # ----------------------------------
    pdf_bytes = None

    try:
        docx_bytes = _render_docx_template_for_request(req, tipo_real, verify_url=verify_url)
    except Exception as e:
        log.exception("Error usando plantilla DOCX para %s #%s: %s", tipo_real, req.id, e)
        docx_bytes = None

    if docx_bytes:
        try:
            class _U:
                def __init__(self, b): self._b = b
                def chunks(self, csize=65536):
                    mv = memoryview(self._b)
                    for i in range(0, len(mv), csize):
                        yield mv[i:i+csize]

            raw_pdf = _office_to_pdf_bytes(
                _U(docx_bytes),
                name_hint=f"{tipo_real}-{req.id}.docx",
            )
        except Exception as e:
            log.exception("Error convirtiendo DOCX→PDF para %s #%s: %s", tipo_real, req.id, e)
            raw_pdf = None

        if raw_pdf:
            try:
                pdf_bytes = _pdf_secure(raw_pdf, user_pwd="", watermark_text="CES · Solo lectura")
            except Exception:
                pdf_bytes = raw_pdf

    # ----------------------------------
    # Fallback: HTML + WeasyPrint (excepto DC3)
    # ----------------------------------
    if pdf_bytes is None and tipo_real != "dc3":
        try:
            tpl, ctx = _render_ctx_and_template(request, req, tipo_real, use_published_if_exists=True)
            html = render_to_string(tpl, ctx, request=request)

            from weasyprint import HTML
            raw_pdf = HTML(string=html, base_url=request.build_absolute_uri("/")).write_pdf()

            pdf_bytes = _pdf_secure(raw_pdf, user_pwd="", watermark_text="CES · Solo lectura")
        except Exception as e:
            log.exception("Error generando PDF HTML fallback para Request %s: %s", req.id, e)
            pdf_bytes = None

    # ----------------------------------
    # Preparar correo
    # ----------------------------------
    extra_msg = (request.POST.get("extra_message") or "").strip()
    extra_mode = (request.POST.get("extra_mode") or "append").strip().lower()
    if extra_mode not in ("append", "full"):
        extra_mode = "append"

    nombre = f"{req.name or ''} {req.lastname or ''}".strip() or "alumno"
    kind_label = "diploma" if tipo_real == "diploma" else "constancia"

    subject = f"Tu {kind_label} – Centro de Estudios Superiores en Negocios y Humanidades"

    # Base TEXT
    base_text = (
        f"Hola {nombre},\n\n"
        f"Tu {kind_label} ha sido publicado.\n"
        f"Puedes verlo aquí:\n{verify_url}\n\n"
        "Centro de Estudios Superiores en Negocios y Humanidades"
    )

    # Base HTML
    base_html = f"""
    <html><body style="font-family:Arial;">
      <div style="text-align:center;margin-bottom:20px;">
        <img src="cid:ceslogo" style="max-width:180px;">
      </div>
      <p>Hola {escape(nombre)},</p>
      <p>Tu {kind_label} ha sido publicado.</p>
      <p>Puedes verlo aquí:</p>
      <p><a href="{verify_url}">Verifica aquí</a></p>
      <p>Centro de Estudios Superiores en Negocios y Humanidades</p>
    </body></html>
    """

    # Aplicar modo FULL o APPEND
    if extra_mode == "full" and extra_msg:
        text_body = extra_msg
        html_body = f"""
        <html><body style="font-family:Arial;">
          <div style="text-align:center;margin-bottom:20px;">
            <img src="cid:ceslogo" style="max-width:180px;">
          </div>
          {escape(extra_msg).replace("\n","<br>")}
          <br><br><a href="{verify_url}">Verifica aquí</a>
        </body></html>
        """
    else:
        text_body = base_text
        if extra_msg:
            text_body += "\n\nMensaje adicional:\n" + extra_msg
        html_body = base_html.replace(
            "</body>",
            f"<p><strong>Mensaje adicional:</strong></p><p>{escape(extra_msg).replace('\n','<br>')}</p></body>"
        ) if extra_msg else base_html

    # ----------------------------------
    # Crear y enviar correo
    # ----------------------------------
    from_email = getattr(settings, "DEFAULT_FROM_EMAIL", None) or getattr(settings, "EMAIL_HOST_USER", None)
    if not from_email:
        return JsonResponse({"ok": False, "error": "Servidor sin remitente configurado."}, status=500)

    email = EmailMultiAlternatives(subject, text_body, from_email, [req.email])
    email.attach_alternative(html_body, "text/html")

    # ---- Adjuntar LOGO con finders ----
    try:
        logo_path = finders.find("img/CES_Logo.png")  # PATH real, no URL
        if logo_path:
            with open(logo_path, "rb") as f:
                logo_data = f.read()
            img = MIMEImage(logo_data)
            img.add_header("Content-ID", "<ceslogo>")
            img.add_header("Content-Disposition", "inline")
            email.attach(img)
        else:
            log.warning("Logo img/CES_Logo.png no fue encontrado en staticfiles.")
    except Exception as e:
        log.exception("No se pudo adjuntar el logo CES: %s", e)

    # Adjuntar PDF (si existe)
    if pdf_bytes:
        filename = f"{tipo_real}-{req.id}.pdf"
        email.attach(filename, pdf_bytes, "application/pdf")

    # Enviar
    try:
        email.send(fail_silently=False)
        emailed = True
    except Exception as e:
        log.exception("Error enviando correo: %s", e)
        emailed = False

    # ----------------------------------
    # Actualizar tracking
    # ----------------------------------
    if emailed:
        req.status = "emailed"
        req.save(update_fields=["status"])

        RequestEvent.objects.update_or_create(
            request=req,
            status="emailed",
            defaults={"note": "Documento enviado por correo."},
        )

        grad = getattr(req, "graduate", None)
        if grad and not grad.sent_at:
            grad.sent_at = timezone.now()
            grad.save(update_fields=["sent_at"])

        # Si es CPROEM → finalizado
        if tipo_real == "cproem":
            req.status = "finalizado"
            req.save(update_fields=["status"])

            RequestEvent.objects.update_or_create(
                request=req,
                status="finalizado",
                defaults={"note": "Constancia CPROEM lista para descarga."},
            )

            if grad and not grad.download_date:
                grad.download_date = timezone.localdate()
                grad.save(update_fields=["download_date"])
    else:
        return JsonResponse(
            {"ok": False, "error": "No se pudo enviar el correo. Revisa las credenciales SMTP."},
            status=500,
        )

    # ----------------------------------
    # Todo correcto
    # ----------------------------------
    return JsonResponse(
        {"ok": True, "verify_url": verify_url, "tipo": tipo_real}
    )


@require_GET
def doc_email_preview(request, req_id):
    """
    Devuelve el mensaje estándar de correo para prellenar el textarea
    cuando el usuario selecciona 'reemplazar'.
    """
    req = get_object_or_404(
        Request.objects.select_related("program", "graduate"),
        pk=req_id,
    )

    tipo_real = "diploma"

    # Generar verify_url sin enviar correo todavía
    doc = _ensure_published_token(req, tipo_real)
    verify_url = request.build_absolute_uri(
        reverse("administracion:verificar_token", args=[doc.token])
    )

    nombre = f"{req.name} {req.lastname}".strip()

    # OJO: aquí ya NO metemos la URL larga, solo el texto
    texto = (
        f"Hola {nombre},\n\n"
        "Tu diploma ha sido publicado.\n"
        "Puedes verlo y verificarlo en el siguiente enlace:\n"
        "Verifica aquí\n\n"
        "Centro de Estudios Superiores en Negocios y Humanidades"
    )

    return JsonResponse({"ok": True, "message": texto})

@require_POST
def doc_confirm(request, req_id):
    import logging
    log = logging.getLogger(__name__)

    # Traemos también el graduate para poder marcar download_date
    req = get_object_or_404(
        Request.objects.select_related("program", "graduate"),
        pk=req_id,
    )

    # --- Tipo solicitado (query) y tipo real ---
    tipo_query = _get_tipo_from_request(request, default="constancia")
    if tipo_query not in ("diploma", "constancia"):
        return JsonResponse({"ok": False, "error": "Tipo inválido."}, status=400)

    tipo_real = "diploma" if tipo_query == "diploma" else constancia_kind_for_request(req)

    # --- Logging ---
    log.info("[doc_confirm] tipo_query=%s → tipo_real=%s (req=%s)", tipo_query, tipo_real, req.id)

    # --- Publicar token / URL de verificación ---
    doc = _ensure_published_token(req, tipo_real)
    verify_url = request.build_absolute_uri(
        reverse("administracion:verificar_token", args=[doc.token])
    )

    # === Solo para CPROEM: marcar como FINALIZADO (círculo 6) ===
    is_cproem = (tipo_real == "cproem")
    if is_cproem:
        # 1) Cambiar estado a 'finalizado'
        if req.status != "finalizado":
            req.status = "finalizado"
            req.save(update_fields=["status"])

        # 2) Registrar evento 'finalizado' una sola vez
        if not RequestEvent.objects.filter(request=req, status="finalizado").exists():
            RequestEvent.objects.create(
                request=req,
                status="finalizado",
                note="Constancia CPROEM confirmada; disponible para descarga.",
            )

        # 3) Marcar fecha de 'descarga' (usamos esto como finalización)
        grad = getattr(req, "graduate", None)
        if grad and not getattr(grad, "download_date", None):
            grad.download_date = timezone.localdate()
            grad.save(update_fields=["download_date"])

    return JsonResponse({
        "ok": True,
        "verify_url": verify_url,
        "tipo": tipo_real,
        "finalizado": is_cproem,
    })


def doc_preview(request, req_id):
    """
    Genera vista previa del documento en PDF.
    Intenta primero DOCX → PDF.
    """
    req = get_object_or_404(
        Request.objects.select_related("program", "graduate"),
        pk=req_id,
    )

    tipo_real = _get_tipo_from_request(request, default="diploma")
    verify_url = request.build_absolute_uri("/preview/")  # no importa, solo para DOCX

    # DOCX primero
    pdf_bytes = None
    try:
        docx_bytes = _render_docx_template_for_request(req, tipo_real, verify_url=verify_url)
        if docx_bytes:
            class _B:
                def __init__(self, b): self.b = b
                def chunks(self, s=65536):
                    mv = memoryview(self.b)
                    for i in range(0, len(mv), s):
                        yield mv[i:i+s]

            pdf_bytes = _office_to_pdf_bytes(_B(docx_bytes),
                                             name_hint=f"preview-{req.id}.docx")
            pdf_bytes = _pdf_secure(pdf_bytes, user_pwd="", watermark_text="CES · SOLO PREVIEW")
    except:
        pdf_bytes = None

    # fallback HTML
    if pdf_bytes is None:
        tpl, ctx = _render_ctx_and_template(request, req, tipo_real)
        html = render_to_string(tpl, ctx, request=request)
        from weasyprint import HTML
        raw = HTML(string=html, base_url=request.build_absolute_uri("/")).write_pdf()
        pdf_bytes = _pdf_secure(raw, user_pwd="", watermark_text="CES · SOLO PREVIEW")

    return HttpResponse(pdf_bytes, content_type="application/pdf")


@require_GET
def doc_download(request, req_id: int):
    """
    Descarga documento del egresado:
      - ?tipo=diploma|constancia
      - ?fmt=pdf|docx|zip
      - (solo CPROEM) ?sig=signed|unsigned
    """
    # imports locales (evitan fallo si libs no instaladas hasta que se usen)
    try:
        from weasyprint import HTML
    except Exception:
        HTML = None

    # --- 1) Obtener registro y params ---
    req = get_object_or_404(Request.objects.select_related("program"), pk=req_id)

    tipo = (request.GET.get("tipo") or "diploma").lower()
    if tipo not in ("diploma", "constancia"):
        return JsonResponse({"ok": False, "error": "Tipo inválido."}, status=400)

    fmt = (request.GET.get("fmt") or "pdf").lower()
    if fmt not in ("pdf", "docx", "zip"):
        return JsonResponse({"ok": False, "error": "Formato inválido."}, status=400)

    # determinar tipo real (dc3 | cproem | diploma)
    tipo_real = "diploma" if tipo == "diploma" else constancia_kind_for_request(req)

    # sig param (solo para cproem)
    sig_param = (request.GET.get("sig") or "").lower()
    if sig_param not in ("", "signed", "unsigned"):
        return JsonResponse({"ok": False, "error": "Parámetro 'sig' inválido."}, status=400)
    signed_flag = None
    if tipo_real == "cproem":
        # default: True (signed) si no se pasa nada o se pasa 'signed'
        signed_flag = (sig_param != "unsigned")

    # intentar publicar token (no crítico)
    try:
        _ensure_published_token(req, tipo_real)
    except Exception:
        log.exception("No se pudo asegurar token publicado para %s #%s", tipo_real, req.id)

    # datos comunes
    program_name = getattr(getattr(req, "program", None), "name", "") or "—"
    full_name = f"{getattr(req, 'name', '') or ''} {getattr(req, 'lastname', '') or ''}".strip()
    start_date = getattr(req, "start_date", None)
    end_date = getattr(req, "end_date", None)

    def _fmt_date(d):
        try:
            return d.strftime("%d/%m/%Y") if d else "—"
        except Exception:
            return str(d) if d else "—"

    # helper: render -> PDF y luego _pdf_secure
    def _pdf_from_html_string(html_str: str) -> bytes:
        if HTML is None:
            raise RuntimeError("WeasyPrint no está instalado.")
        base_url = request.build_absolute_uri("/")
        raw_pdf = HTML(string=html_str, base_url=base_url).write_pdf()
        secured = _pdf_secure(
            raw_pdf,
            user_pwd="",
            watermark_text="CES · Solo lectura",
        )
        return secured

    # Construye contexto para plantilla CPROEM (firma incluida solo si include_signature=True)
    def _build_cproem_render_context(req, graduate, include_signature: bool):
        nombre = (f"{req.name or ''} {req.lastname or ''}".strip()) or getattr(graduate, "nombre", "") or "—"
        programa = getattr(getattr(req, "program", None), "name", "") or "—"

        # vigencias: preferir datos del graduate si existen
        vigencia_inicio = getattr(graduate, "validity_start", None)
        vigencia_termino = getattr(graduate, "validity_end", None)

        # fallback fechas petición
        inicio = getattr(req, "start_date", None)
        fin = getattr(req, "end_date", None)

        # QR: usamos el token publicado si existe; si no, generamos uno nuevo
        published = _get_existing_token(req, "cproem")
        token = published.token if published else _random_hex_25()
        verify_url = request.build_absolute_uri(
            reverse("administracion:verificar_token", args=[token])
        )
        # PNG embebido, igual que en diplomas/DC3
        qr_url = _qr_png_data_url(verify_url, box_size=12)

        # Firma: solo si include_signature -> construimos URL a static
        signature_url = ""
        if include_signature:
            SIGNATURE_STATIC_PATH = "img/signatures/firma_cproem.png"  # ajusta si tu firma está en otra ruta
            try:
                signature_url = request.build_absolute_uri(static(SIGNATURE_STATIC_PATH))
            except Exception:
                signature_url = ""

        return {
            "nombre": nombre,
            "programa": programa,
            "vigencia_inicio": vigencia_inicio,
            "vigencia_termino": vigencia_termino,
            "inicio": inicio,
            "fin": fin,
            "curp": getattr(req, "curp", "") or "",
            "qr_url": qr_url,
            "signature_url": signature_url,
        }

    # Generadores PDF/DOCX (HTML fallback)
    def _make_pdf_bytes_for_cproem_html(graduate, signed: bool) -> bytes:
        """
        Versión HTML (plantillas antiguas) – ahora solo como fallback si no hay DOCX.
        """
        try:
            # Si signed -> plantilla digital (contiene la firma), si unsigned -> plantilla sin firma
            tpl_signed = "administracion/pdf_constancia_cproem_digital.html"
            tpl_unsigned = "administracion/pdf_constancia_cproem.html"

            tpl_name = tpl_signed if signed else tpl_unsigned
            ctx = _build_cproem_render_context(req, graduate, include_signature=signed)
            html = render_to_string(tpl_name, ctx, request=request)
            return _pdf_from_html_string(html)
        except Exception as e:
            raise RuntimeError(f"Error generando documento: {e}")

    def _make_pdf_bytes_generic_html() -> bytes:
        """
        Fallback genérico HTML → PDF para tipos sin DOCX.
        """
        try:
            tpl, ctx = _render_ctx_and_template(request, req, tipo_real, use_published_if_exists=True)
            html = render_to_string(tpl, ctx, request=request)
            return _pdf_from_html_string(html)
        except Exception as e:
            raise RuntimeError(f"Error generando documento: {e}")

    def _make_docx_bytes_generic(include_cproem_note=False, signed=None):
        # 1) Intentar plantilla DOCX del programa (docx con variables)
        try:
            # Para el DOCX queremos también URL de verificación
            doc_token = _ensure_published_token(req, tipo_real)
            verify_url = request.build_absolute_uri(
                reverse("administracion:verificar_token", args=[doc_token.token])
            )

            rendered = _render_docx_template_for_request(
                req,
                tipo_real,
                verify_url=verify_url,
            )
            if rendered:
                return rendered
        except Exception as e:
            # Si falla la plantilla, logueamos y caemos al genérico
            log.exception(
                "Error usando plantilla DOCX para %s #%s: %s",
                tipo_real,
                req.id,
                e,
            )

        # 2) Fallback: DOCX genérico como antes
        try:
            doc = Document()
        except Exception as e:
            raise RuntimeError(f"python-docx no está instalado: {e}")

        title = (
            f"{tipo_real.upper()} - {program_name}"
            if tipo_real != "cproem"
            else f"CONSTANCIA CPROEM - {program_name}"
        )
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Nombre: {full_name or '—'}")
        doc.add_paragraph(f"Programa: {program_name}")
        doc.add_paragraph(f"Inicio: {_fmt_date(start_date)}")
        doc.add_paragraph(f"Fin: {_fmt_date(end_date)}")

        curp = getattr(req, "curp", None)
        rfc = getattr(req, "rfc", None)
        job = getattr(req, "job_title", None)
        giro = getattr(req, "industry", None)
        biz = getattr(req, "business_name", None)

        if curp:
            doc.add_paragraph(f"CURP: {curp}")
        if rfc:
            doc.add_paragraph(f"RFC: {rfc}")
        if job:
            doc.add_paragraph(f"Puesto: {job}")
        if giro:
            doc.add_paragraph(f"Giro: {giro}")
        if biz:
            doc.add_paragraph(f"Razón social: {biz}")

        if include_cproem_note:
            doc.add_paragraph("")
            if signed is False:
                doc.add_paragraph("Versión sin firma (sin imagen de firma).")
            else:
                doc.add_paragraph("Versión con firma (contiene imagen de firma).")

        doc.add_paragraph("")
        doc.add_paragraph("Documento generado automáticamente por CES System.")

        b = io.BytesIO()
        doc.save(b)
        b.seek(0)
        return b.getvalue()

    # --- 4) Ramas por formato solicitado ---

    # =========================
    #    C P R O E M
    # =========================
    if tipo_real == "cproem":
        # localizar graduate existente
        grad = getattr(req, "graduate", None)
        if not grad:
            grad = Graduate.objects.filter(request=req).first()

        # si no existe, lo creamos automáticamente con datos mínimos de la Request
        if not grad:
            grad, _ = Graduate.objects.get_or_create(
                request=req,
                defaults={
                    "name": getattr(req, "name", "") or "",
                    "lastname": getattr(req, "lastname", "") or "",
                    "email": getattr(req, "email", "") or "",
                    "curp": getattr(req, "curp", "") or "",
                    "job_title": getattr(req, "job_title", "") or "",
                    "industry": getattr(req, "industry", "") or "",
                    "business_name": getattr(req, "business_name", "") or "",
                    "url": getattr(req, "url", "") or "",
                    "validity_start": getattr(req, "start_date", None),
                    "validity_end": getattr(req, "end_date", None),
                },
            )

        # --------- PDF (CPROEM) ---------
        if fmt == "pdf":
            pdf_bytes = None

            # 1) Intentar DOCX con variables -> PDF
            try:
                doc_token = _ensure_published_token(req, tipo_real)
                verify_url = request.build_absolute_uri(
                    reverse("administracion:verificar_token", args=[doc_token.token])
                )

                docx_bytes = _render_docx_template_for_request(
                    req,
                    "cproem",
                    verify_url=verify_url,
                )
            except Exception as e:
                log.exception(
                    "Error usando plantilla DOCX para CPROEM #%s: %s",
                    req.id,
                    e,
                )
                docx_bytes = None

            if docx_bytes:
                try:
                    class _U:
                        def __init__(self, b): self._b = b
                        def chunks(self, csize=64 * 1024):
                            mv = memoryview(self._b)
                            for i in range(0, len(mv), csize):
                                yield mv[i:i + csize]

                    raw_pdf = _office_to_pdf_bytes(
                        _U(docx_bytes),
                        name_hint=f"{tipo_real}-{req.id}.docx",
                    )
                    pdf_bytes = _pdf_secure(
                        raw_pdf,
                        user_pwd="",
                        watermark_text="CES · Solo lectura",
                    )
                except Exception as e:
                    log.exception(
                        "Error DOCX→PDF CPROEM #%s: %s",
                        req.id,
                        e,
                    )
                    pdf_bytes = None

            # 2) Fallback HTML (plantillas CPROEM antiguas con firma)
            if pdf_bytes is None:
                try:
                    signed = signed_flag if signed_flag is not None else True
                    pdf_bytes = _make_pdf_bytes_for_cproem_html(grad, signed=signed)
                except Exception as e:
                    log.exception(
                        "Fallo generando documento CPROEM #%s (fmt=pdf sig=%s)",
                        req.id,
                        signed_flag,
                    )
                    return JsonResponse({"ok": False, "error": str(e)}, status=500)

            return FileResponse(
                io.BytesIO(pdf_bytes),
                content_type="application/pdf",
                filename=f"{tipo_real}-{req.id}.pdf",
            )

        # --------- DOCX (CPROEM) ---------
        if fmt == "docx":
            try:
                # Usa plantilla DOCX del programa si existe (a través de _make_docx_bytes_generic)
                docx_bytes = _make_docx_bytes_generic(
                    include_cproem_note=True,
                    signed=(signed_flag if signed_flag is not None else True),
                )
                return FileResponse(
                    io.BytesIO(docx_bytes),
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    filename=f"{tipo_real}-{req.id}.docx",
                )
            except Exception as e:
                log.exception("Fallo generando DOCX CPROEM #%s", req.id)
                return JsonResponse({"ok": False, "error": str(e)}, status=500)

        # --------- ZIP (CPROEM) ---------
        if fmt == "zip":
            mem = io.BytesIO()
            readme_lines = []
            try:
                with zipfile.ZipFile(mem, "w", zipfile.ZIP_DEFLATED) as zf:
                    # signed pdf (DOCX→PDF si hay, si no HTML)
                    try:
                        # DOCX→PDF
                        try:
                            doc_token = _ensure_published_token(req, tipo_real)
                            verify_url = request.build_absolute_uri(
                                reverse("administracion:verificar_token", args=[doc_token.token])
                            )
                            docx_bytes = _render_docx_template_for_request(
                                req, "cproem", verify_url=verify_url
                            )
                        except Exception:
                            docx_bytes = None

                        if docx_bytes:
                            class _U2:
                                def __init__(self, b): self._b = b
                                def chunks(self, csize=64 * 1024):
                                    mv = memoryview(self._b)
                                    for i in range(0, len(mv), csize):
                                        yield mv[i:i + csize]

                            raw_signed = _office_to_pdf_bytes(
                                _U2(docx_bytes),
                                name_hint=f"{tipo_real}-{req.id}.docx",
                            )
                            signed_pdf = _pdf_secure(
                                raw_signed,
                                user_pwd="",
                                watermark_text="CES · Solo lectura",
                            )
                        else:
                            # fallback HTML firmado
                            signed_pdf = _make_pdf_bytes_for_cproem_html(grad, signed=True)

                        zf.writestr(
                            f"{tipo_real}-{req.id}-signed.pdf",
                            signed_pdf,
                        )
                    except Exception as e:
                        msg = f"[PDF signed] fallo: {e}"
                        log.exception(msg)
                        readme_lines.append(msg)

                    # unsigned pdf (para zip, seguimos con HTML sin firma)
                    try:
                        unsigned_pdf = _make_pdf_bytes_for_cproem_html(grad, signed=False)
                        zf.writestr(
                            f"{tipo_real}-{req.id}-unsigned.pdf",
                            unsigned_pdf,
                        )
                    except Exception as e:
                        msg = f"[PDF unsigned] fallo: {e}"
                        log.exception(msg)
                        readme_lines.append(msg)

                    # docx (plantilla con variables si existe, si no genérico)
                    try:
                        zf.writestr(
                            f"{tipo_real}-{req.id}.docx",
                            _make_docx_bytes_generic(
                                include_cproem_note=True, signed=True
                            ),
                        )
                    except Exception as e:
                        msg = f"[DOCX] fallo: {e}"
                        log.exception(msg)
                        readme_lines.append(msg)

                    if readme_lines:
                        zf.writestr(
                            "README.txt",
                            "Algunos archivos no se pudieron generar:\n\n"
                            + "\n".join(readme_lines),
                        )

                mem.seek(0)
                return FileResponse(
                    mem,
                    content_type="application/zip",
                    filename=f"{tipo_real}-{req.id}.zip",
                )
            except Exception as e:
                log.exception("Fallo generando ZIP CPROEM #%s", req.id)
                return JsonResponse(
                    {"ok": False, "error": f"Error generando ZIP: {e}"},
                    status=500,
                )

    # =========================
    #  D I P L O M A / D C 3
    # =========================
    if fmt == "pdf":
        # 1) Intentar generar PDF a partir de una plantilla DOCX con variables (si existe)
        docx_bytes = None
        try:
            doc_token = _ensure_published_token(req, tipo_real)
            verify_url = request.build_absolute_uri(
                reverse("administracion:verificar_token", args=[doc_token.token])
            )

            docx_bytes = _render_docx_template_for_request(
                req,
                tipo_real,
                verify_url=verify_url,
            )
        except Exception as e:
            log.exception(
                "Error usando plantilla DOCX para PDF %s #%s: %s",
                tipo_real,
                req.id,
                e,
            )
            docx_bytes = None

        if docx_bytes:
            # convertir DOCX -> PDF usando LibreOffice (headless), si está disponible
            pdf_bytes = None
            try:
                class _U:
                    def __init__(self, b): self._b = b
                    def chunks(self, csize=64 * 1024):
                        mv = memoryview(self._b)
                        for i in range(0, len(mv), csize):
                            yield mv[i:i + csize]

                pdf_bytes = _office_to_pdf_bytes(
                    _U(docx_bytes),
                    name_hint=f"{tipo_real}-{req.id}.docx",
                )
            except Exception:
                pdf_bytes = None

            if pdf_bytes:
                # endurecer PDF igual que en el flujo HTML (raster + watermark + permisos)
                try:
                    secured_pdf = _pdf_secure(
                        pdf_bytes,
                        user_pwd="",
                        watermark_text="CES · Solo lectura",
                    )
                except Exception:
                    secured_pdf = pdf_bytes

                return FileResponse(
                    io.BytesIO(secured_pdf),
                    content_type="application/pdf",
                    filename=f"{tipo_real}-{req.id}.pdf",
                )

        # 2) Fallback: comportamiento previo (HTML + imagen de fondo con WeasyPrint)
        try:
            pdf_bytes = _make_pdf_bytes_generic_html()
            return FileResponse(
                io.BytesIO(pdf_bytes),
                content_type="application/pdf",
                filename=f"{tipo_real}-{req.id}.pdf",
            )
        except Exception as e:
            log.exception("Fallo generando PDF")
            return JsonResponse(
                {"ok": False, "error": f"Error generando PDF: {e}"},
                status=500,
            )

    if fmt == "docx":
        try:
            docx_bytes = _make_docx_bytes_generic(include_cproem_note=False)
            return FileResponse(
                io.BytesIO(docx_bytes),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"{tipo_real}-{req.id}.docx",
            )
        except Exception as e:
            log.exception("Fallo generando DOCX")
            return JsonResponse(
                {"ok": False, "error": f"Error generando DOCX: {e}"},
                status=500,
            )

    # zip genérico: PDF + DOCX (diploma/dc3)
    readme_lines = []
    mem = io.BytesIO()
    with zipfile.ZipFile(mem, "w", zipfile.ZIP_DEFLATED) as zf:
        try:
            zf.writestr(
                f"{tipo_real}-{req.id}.docx",
                _make_docx_bytes_generic(include_cproem_note=False),
            )
        except Exception as e:
            msg = f"[DOCX] No se pudo generar: {e}"
            log.exception(msg)
            readme_lines.append(msg)
        try:
            zf.writestr(
                f"{tipo_real}-{req.id}.pdf",
                _make_pdf_bytes_generic_html(),
            )
        except Exception as e:
            msg = f"[PDF] No se pudo generar: {e}"
            log.exception(msg)
            readme_lines.append(msg)

        if readme_lines:
            zf.writestr(
                "README.txt",
                "Algunos archivos no se pudieron generar:\n\n"
                + "\n".join(readme_lines),
            )

    mem.seek(0)
    return FileResponse(
        mem,
        content_type="application/zip",
        filename=f"{tipo_real}-{req.id}.zip",
    )


# ========================= Verificación =========================

def verificar_token(request, token: str):
    dt = get_object_or_404(DocToken, token=token, is_active=True)
    req = dt.request

    grad = getattr(req, 'graduate', None)
    if grad:
        inicio = getattr(grad, 'validity_start', None)
        fin = getattr(grad, 'validity_end', None)
    else:
        inicio = getattr(req, 'start_date', None)
        fin = getattr(req, 'end_date', None)

    ctx = {
        "token": token,
        "tipo": dt.tipo,
        "alumno": f"{getattr(req,'name','')} {getattr(req,'lastname','')}".strip(),
        "programa": getattr(getattr(req, "program", None), "name", "—"),
        "inicio": inicio,
        "fin": fin,
    }
    return render(request, "administracion/verify_result.html", ctx)

# ====================== Tokens helpers ======================

def _ensure_published_token(req: Request, tipo: str) -> DocToken:
    try:
        return DocToken.objects.get(request=req, tipo=tipo, is_active=True)
    except DocToken.DoesNotExist:
        token = _random_hex_25()
        return DocToken.objects.create(request=req, tipo=tipo, token=token, is_active=True)

def _get_existing_token(req: Request, tipo: str):
    try:
        return DocToken.objects.get(request=req, tipo=tipo, is_active=True)
    except DocToken.DoesNotExist:
        return None


# ====================== Render helpers (ctx/plantilla/PDF/DOCX) ======================

def _render_ctx_and_template(request, req, tipo_real: str, *, use_published_if_exists: bool):
    # Nombre completo y programa
    nombre = f"{getattr(req,'name','')} {getattr(req,'lastname','')}".strip()

    program_obj = getattr(req, "program", None)
    programa = getattr(program_obj, "name", "") or ""
    program_name_lower = programa.lower()

    # Fechas del curso: primero Graduate, luego Request
    grad = getattr(req, 'graduate', None)
    if grad:
        inicio = getattr(grad, 'validity_start', None)
        fin = getattr(grad, 'validity_end', None)
    else:
        inicio = getattr(req, 'start_date', None)
        fin = getattr(req, 'end_date', None)

    # Token de verificación
    if use_published_if_exists:
        published = _get_existing_token(req, tipo_real)
        token = published.token if published else _random_hex_25()
    else:
        token = _random_hex_25()

    verify_url = request.build_absolute_uri(
        reverse("administracion:verificar_token", args=[token])
    )
    qr_url = _qr_png_data_url(verify_url, box_size=12)

    # ======================================================
    #  Diplomas
    # ======================================================
    if tipo_real == "diploma":
        # Fondo según programa abreviado
        bg_static, show_program_text = _diploma_background_for_request(req)

        # Si bg_static es None, simplemente no usamos fondo
        bg_url = static(bg_static) if bg_static else ""

        tpl = "administracion/pdf_diploma.html"
        ctx = {
            "folio": f"DIP-{req.id:06d}",
            "nombre": nombre,
            "programa": programa,
            "inicio": inicio,
            "fin": fin,
            "qr_url": qr_url,
            "bg_url": bg_url,
            "show_program_text": show_program_text,
        }

    # ======================================================
    #  Constancias (DC-3 y CPROEM)
    # ======================================================
    else:
        tpl = (
            "administracion/pdf_constancia_dc3.html"
            if tipo_real == "dc3"
            else "administracion/pdf_constancia_cproem.html"
        )

        ctx = {
            "folio": f"CON-{req.id:06d}",
            "nombre": nombre,
            "programa": programa,
            "inicio": inicio,
            "fin": fin,
            "curp": getattr(req, "curp", None),
            "rfc": getattr(req, "rfc", None),
            "puesto": getattr(req, "job_title", None),
            "giro": getattr(req, "industry", None),
            "razon_social": getattr(req, "business_name", None),
            "horas": getattr(req, "hours", None),
            "qr_url": qr_url,
        }

    return tpl, ctx


def _diploma_background_for_request(req):
    """
    Elige automáticamente el fondo del diploma según el programa abreviado.

    Regla:
      - Toma el código del programa (abbreviation / programa / code).
      - Intenta primero nombres "bonitos" estándar:
          1) static/img/diplomas/<CODIGO>.png
          2) static/img/diplomas/<CODIGO>N.png
      - Si no existen, busca en static/img/diplomas cualquier PNG que
        contenga <CODIGO> en el nombre (por ejemplo: 'Diploma_DGE_2025.png').
      - Si aún así no hay nada, devuelve el primer candidato aunque no exista;
        en la práctica verás la hoja en blanco (señal de que falta el PNG).
    """
    program_obj = getattr(req, "program", None)

    code = ""
    if program_obj:
        code = (
            getattr(program_obj, "abbreviation", None)
            or getattr(program_obj, "programa", None)
            or getattr(program_obj, "code", None)
            or ""
        )

    code = (code or "").upper().strip()

    # Si no hay código, ni lo intentamos
    if not code:
        return None, True

    # 1) Candidatos estándar
    candidates = [
        f"img/diplomas/{code}.png",
        f"img/diplomas/{code}N.png",
    ]

    # ¿Alguno existe en static?
    for rel_path in candidates:
        if finders.find(rel_path):
            # Para fondos específicos normalmente el título ya va impreso
            return rel_path, False

    # 2) Búsqueda flexible en static/img/diplomas:
    #    cualquier PNG cuyo nombre contenga la abreviatura.
    try:
        static_dir = Path(settings.BASE_DIR) / "static" / "img" / "diplomas"
        if static_dir.exists():
            for fname in os.listdir(static_dir):
                if not fname.lower().endswith(".png"):
                    continue
                if code in fname.upper():
                    # Ej.: 'Diploma_DGE_2025.png' -> 'img/diplomas/Diploma_DGE_2025.png'
                    return f"img/diplomas/{fname}", False
    except Exception:
        # fallback silencioso, no queremos romper la generación del PDF
        pass

    # 3) Ninguno encontrado: devolvemos el primero como fallback.
    # En la práctica se verá blanco, lo que te avisa que falta el PNG real.
    return candidates[0], True

def _dc3_background_for_request(req):
    """
    Devuelve rutas *relativas* a static para el frente y reverso de la DC-3,
    de acuerdo al programa abreviado.

    Archivos esperados: STATIC/img/DC3/DC3_<ABREVIATURA>_F.png y _R.png
    Ejemplo: DC3_DAP_F.png, DC3_DAP_R.png
    """
    program = getattr(req, "program", None)

    abbr = ""
    if program is not None:
        abbr = (
            getattr(program, "abbreviation", None)
            or getattr(program, "code", None)
            or ""
        )

    abbr = (abbr or "").upper().strip()

    if not abbr:
        # fondo genérico si no hay abreviatura
        return ("img/DC3/DC3_GENERIC_F.png", "img/DC3/DC3_GENERIC_R.png")

    front_rel = f"img/DC3/DC3_{abbr}_F.png"
    back_rel  = f"img/DC3/DC3_{abbr}_R.png"
    return front_rel, back_rel

def _render_ctx_and_template(request, req, tipo_real: str, *, use_published_if_exists: bool):
    # Nombre completo "normal"
    nombre = f"{getattr(req, 'name', '')} {getattr(req, 'lastname', '')}".strip()

    program_obj = getattr(req, "program", None)
    programa = getattr(program_obj, "name", "") or ""
    program_name_lower = programa.lower()

    # Fechas del curso: primero Graduate, luego Request
    grad = getattr(req, 'graduate', None)
    if grad:
        inicio = getattr(grad, 'validity_start', None)
        fin = getattr(grad, 'validity_end', None)
    else:
        inicio = getattr(req, 'start_date', None)
        fin = getattr(req, 'end_date', None)

    # 🔹 HORAS: primero Graduate, si no hay dejamos None (el template puede usar default)
    horas_val = getattr(grad, "hours", None) if grad else None

    # Token de verificación
    if use_published_if_exists:
        published = _get_existing_token(req, tipo_real)
        token = published.token if published else _random_hex_25()
    else:
        token = _random_hex_25()

    verify_url = request.build_absolute_uri(
        reverse("administracion:verificar_token", args=[token])
    )
    qr_url = _qr_png_data_url(verify_url, box_size=12)

    # ======================================================
    #  D I P L O M A S
    # ======================================================
    if tipo_real == "diploma":
        # Fondo según programa (tu helper actual)
        bg_static, show_program_text = _diploma_background_for_request(req)
        bg_url = static(bg_static) if bg_static else ""

        tpl = "administracion/pdf_diploma.html"
        ctx = {
            "folio": f"DIP-{req.id:06d}",
            "nombre": nombre,
            "programa": programa,
            "inicio": inicio,
            "fin": fin,
            "qr_url": qr_url,
            "bg_url": bg_url,
            "show_program_text": show_program_text,
            "horas": horas_val,  # 👈 NUEVO: intensidad para el diploma
        }

    # ======================================================
    #  C O N S T A N C I A S  (DC-3 y CPROEM)
    # ======================================================
    else:
        if tipo_real == "dc3":
            # Nombre especial para DC3: "APELLIDOS NOMBRE"
            first = getattr(req, "name", "") or ""
            last = getattr(req, "lastname", "") or ""
            dc3_nombre = f"{last} {first}".strip()

            # Fondos específicos por programa (F/R)
            front_rel, back_rel = _dc3_background_for_request(req)
            dc3_front_url = static(front_rel) if front_rel else ""
            dc3_back_url  = static(back_rel) if back_rel else ""

            tpl = "administracion/pdf_constancia_dc3.html"
            ctx = {
                "folio": f"CON-{req.id:06d}",
                "nombre": nombre,                     # por si lo necesitas en otro lado
                "dc3_nombre": dc3_nombre or nombre,   # ESTE es el que usa el HTML
                "programa": programa,
                "inicio": inicio,
                "fin": fin,
                "curp": getattr(req, "curp", None),
                "rfc": getattr(req, "rfc", None),
                "puesto": getattr(req, "job_title", None),
                "giro": getattr(req, "industry", None),
                "razon_social": getattr(req, "business_name", None),
                "horas": horas_val,                   # 👈 si la quieres también en DC3
                "qr_url": qr_url,
                "dc3_front_url": dc3_front_url,
                "dc3_back_url": dc3_back_url,
            }
        else:
            # CPROEM
            tpl = "administracion/pdf_constancia_cproem.html"
            ctx = {
                "folio": f"CON-{req.id:06d}",
                "nombre": nombre,
                "programa": programa,
                "inicio": inicio,
                "fin": fin,
                "curp": getattr(req, "curp", None),
                "rfc": getattr(req, "rfc", None),
                "puesto": getattr(req, "job_title", None),
                "giro": getattr(req, "industry", None),
                "razon_social": getattr(req, "business_name", None),
                "horas": horas_val,                   # 👈 igual aquí si algún día la usas
                "qr_url": qr_url,
            }

    return tpl, ctx


# ==================== Seguridad PDF ====================

def _pdf_secure(pdf_bytes: bytes, *, user_pwd: str = "", watermark_text: Optional[str] = None) -> bytes:
    """
    Endurece el PDF:
      1) Rasteriza cada página (PyMuPDF) → dificulta edición/extracción.
      2) Aplica marca de agua diagonal (opcional) con Pillow.
      3) Cifra con pikepdf, bloqueando copiar/modificar/imprimir.
    Si alguna lib falta, hace "best effort" con lo disponible.
    """
    working_pdf = pdf_bytes

    # 1) Rasterización y marca de agua (opcional)
    try:
        import fitz  # PyMuPDF
        from PIL import Image, ImageDraw, ImageFont

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        mat = fitz.Matrix(2, 2)  # ~144 dpi
        frames = []
        for i in range(doc.page_count):
            pix = doc.load_page(i).get_pixmap(matrix=mat, alpha=False)
            im = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

            if watermark_text:
                W, H = im.size
                draw = ImageDraw.Draw(im)
                try:
                    font = ImageFont.load_default()
                except Exception:
                    font = None

                # Capa temporal RGBA para simular transparencia
                wm = Image.new("RGBA", (W, H), (0, 0, 0, 0))
                wmdraw = ImageDraw.Draw(wm)

                text = str(watermark_text)
                # Repetir texto diagonalmente
                step = max(240, int(min(W, H) / 4))
                angle = -30
                # Dibuja líneas de texto antes de rotar, luego rota la capa
                for y in range(0, H + step, step):
                    for x in range(-W, W, step):
                        wmdraw.text((x, y), text, fill=(180, 180, 180, 70), font=font)

                wm = wm.rotate(angle, expand=1)
                # Centrar la marca rotada sobre el lienzo
                x0 = (W - wm.width) // 2
                y0 = (H - wm.height) // 2
                im = im.convert("RGBA")
                im.alpha_composite(wm, (x0, y0))
                im = im.convert("RGB")

            frames.append(im)

        doc.close()
        if frames:
            out = io.BytesIO()
            frames[0].save(out, format="PDF", save_all=True, append_images=frames[1:])
            working_pdf = out.getvalue()
    except Exception:
        # Si falla PyMuPDF/Pillow, seguimos con el original
        pass

    # 2) Cifrado + permisos con pikepdf
    try:
        import pikepdf, inspect

        desired = {
            "print_lowres": False,
            "print_highres": False,
            "extract": False,       # copiar/extraer
            "modify": False,
            "annotate": False,
            "form": False,
            "assemble": False,
            "accessibility": False, # pon True si requieres accesibilidad
        }
        perm_params = set(inspect.signature(pikepdf.Permissions).parameters.keys())
        perms = pikepdf.Permissions(**{k: v for k, v in desired.items() if k in perm_params})

        enc = pikepdf.Encryption(
            user=(user_pwd or ""),          # "" = abre sin pedir pass
            owner=secrets.token_hex(16),    # owner aleatorio
            R=6,                            # AES-256
            allow=perms
        )

        out = io.BytesIO()
        with pikepdf.open(io.BytesIO(working_pdf)) as pdf:
            pdf.save(out, encryption=enc)
        return out.getvalue()
    except Exception:
        # Si no hay pikepdf, devolvemos lo rasterizado/marcado
        return working_pdf
# ================== Catálogo simple ==================

def plantillas_catalogo(request):
    return render(request, "administracion/plantillas_catalogo.html")


# ================== Tablero de solicitudes ==================

class RequestsBoardView(LoginRequiredMixin, TemplateView):
    template_name = "administracion/solicitudes.html"
    login_url = "administracion:login"

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)

        # Filtros por querystring (compat: 'estado' o 'status')
        q = (self.request.GET.get("q") or "").strip()
        estado = (
            self.request.GET.get("estado")
            or self.request.GET.get("status")
            or "pending"
        ).strip() or "pending"
        program_id = (self.request.GET.get("program") or "").strip()

        # Base queryset
        base = Request.objects.select_related("program")

        if q:
            base = base.filter(
                Q(email__icontains=q)
                | Q(name__icontains=q)
                | Q(lastname__icontains=q)
                | Q(program__name__icontains=q)
            )

        if program_id:
            base = base.filter(program_id=program_id)

        # Orden por estado
        if estado == "pending":
            qs = base.filter(status="pending").order_by("sent_at")
        elif estado in {"accepted", "rejected", "generating", "review", "finalizado"}:
            qs = base.filter(status=estado).order_by("-sent_at")
        else:
            qs = base.order_by("-sent_at")

        # Resolver nombre de programa desde default (ces_db)
        qs_list = list(qs)
        missing_ids = {
            r.program_id
            for r in qs_list
            if getattr(r, "program_id", None)
            and (
                getattr(r, "program", None) is None
                or not getattr(getattr(r, "program", None), "name", None)
            )
        }

        sim_map = {}
        if missing_ids:
            sim_programs = ProgramSim.objects.filter(
                id__in=missing_ids
            ).only("id", "name")
            sim_map = {p.id: getattr(p, "name", "") for p in sim_programs}

        for r in qs_list:
            name = None
            if getattr(r, "program", None) and getattr(r.program, "name", None):
                name = r.program.name
            elif r.program_id:
                name = sim_map.get(r.program_id)
            setattr(r, "program_name", name or "Sin programa")

        items = qs_list

        # Contadores por estado
        all_counts = Request.objects.values("status").annotate(total=Count("id"))
        counts = {c["status"]: c["total"] for c in all_counts}

        # Catálogo de programas (default)
        programas = list(
            ProgramSim.objects.all()
            .order_by("name")
            .values("id", "name", "abbreviation")
        )

        # pestañas/estados disponibles (ya incluye FINALIZADO)
        statuses = [
            ("pending", "Pendiente"),
            ("review", "Revisión"),
            ("accepted", "Aprobada"),
            ("rejected", "Rechazada"),
            ("generating", "Generando"),
            ("finalizado", "Finalizado"),
        ]

        ctx.update(
            {
                "q": q,
                "estado": estado,
                "status": estado,
                "program_id": program_id,
                "items": items,
                "programas": programas,
                "counts": counts,
                "count_pending": counts.get("pending", 0),
                "statuses": statuses,
                "now": timezone.now(),
            }
        )

        ctx["csrf_token"] = get_token(self.request)
        ctx["requests_update_url"] = reverse(
            "administracion:request_update_status"
        )
        return ctx


@login_required
@require_POST
def request_update_status(request):
    """
    Actualiza el status de una solicitud desde el tablero (AJAX).
    Ahora cualquier usuario autenticado y activo puede hacerlo.
    """
    me = request.user
    if not me.is_active:
        return JsonResponse({"ok": False, "msg": "Usuario inactivo."}, status=403)

    rid = request.POST.get("id")
    action = (request.POST.get("action") or "").strip()
    reason = (request.POST.get("reason") or "").strip()

    STATUS_MAP = {
        "to_review":  "review",
        "approve":    "accepted",
        "reject":     "rejected",
        "generating": "generating",
    }

    if not rid or action not in STATUS_MAP:
        return JsonResponse({"ok": False, "msg": "Parámetros inválidos."}, status=400)

    try:
        req_obj = Request.objects.get(pk=rid)
    except Request.DoesNotExist:
        return JsonResponse({"ok": False, "msg": "Solicitud no encontrada."}, status=404)

    new_status = STATUS_MAP[action]

    if new_status == "rejected" and not reason:
        return JsonResponse({"ok": False, "msg": "Falta el motivo de rechazo."}, status=400)

    # Actualizar estatus + motivo (si aplica)
    if new_status == "rejected":
        req_obj.status_reason = reason
        req_obj.status = new_status
        req_obj.save(update_fields=["status", "status_reason"])
    else:
        req_obj.status = new_status
        req_obj.save(update_fields=["status"])

    # Registrar evento de la solicitud (bitácora)
    if not RequestEvent.objects.filter(request=req_obj, status=new_status).exists():
        RequestEvent.objects.create(
            request=req_obj,
            status=new_status,
            note=(reason or None),
        )

    return JsonResponse({"ok": True, "new_status": new_status})
# ===================== (NUEVO) Guardar edición inline =====================

@login_required
@require_http_methods(["POST"])
def egresado_update_inline(request, req_id: int):
    """Actualiza campos básicos del Request desde Egresados (edición inline)."""
    if not _admin_allowed(request.user):
        return JsonResponse({"ok": False, "msg": "No autorizado."}, status=403)

    req_obj = get_object_or_404(Request, pk=req_id)

    # Import local para no tocar cabecera
    from django.utils.dateparse import parse_date

    allowed = ("name","lastname","curp","rfc","job_title","industry","business_name","start_date","end_date")
    changed = {}

    for f in allowed:
        val = (request.POST.get(f) or "").strip()
        if f in ("start_date", "end_date"):
            new_val = parse_date(val) if val else None
            if val and new_val is None:
                return JsonResponse({"ok": False, "msg": f"Fecha inválida en {f} (use YYYY-MM-DD)."}, status=400)
            old_val = getattr(req_obj, f, None)
            if old_val != new_val:
                setattr(req_obj, f, new_val)
                changed[f] = new_val.isoformat() if new_val else ""
        else:
            old_text = getattr(req_obj, f, "") or ""
            if val != old_text:
                setattr(req_obj, f, val)
                changed[f] = val

    if changed:
        req_obj.save(update_fields=list(changed.keys()))
        # Registrar un evento suave (opcional, no afecta flujos)
        if not RequestEvent.objects.filter(request=req_obj, status="review").exists():
            RequestEvent.objects.create(request=req_obj, status="review", note="Edición de datos (admin)")

    return JsonResponse({"ok": True, "changed": changed})


# ================== PLANTILLAS PREMIUM ==================

def plantillas_admin(request):
    q = request.GET.get("q", "").strip()

    # Plantillas de diseño (imágenes / fondos) — lo que ya tenías
    qs = DesignTemplate.objects.all().order_by("-updated_at")
    if q:
        qs = qs.filter(title__icontains=q)

    # NUEVO: plantillas DOCX con variables
    docx_qs = DocxTemplate.objects.all().order_by("-created_at")
    if q:
        docx_qs = docx_qs.filter(name__icontains=q)

    return render(
        request,
        "administracion/plantillas_admin.html",
        {
            "items": qs,
            "docx_items": docx_qs,  # <-- nuevo contexto
            "q": q,
        },
    )

@login_required
@require_http_methods(["GET", "POST"])
def plantilla_create(request):
    if request.method == "POST":
        title = request.POST.get("title") or "Nueva plantilla"
        kind = (request.POST.get("kind") or DesignTemplate.DESIGN).lower()
        raw = (request.POST.get("json") or "").strip()

        if not raw:
            data = {"pages": [{"width": 1920, "height": 1080, "background": "#ffffff", "layers": []}]}
        else:
            try:
                data = json.loads(raw)
                pages = data.get("pages") or []
                if not pages:
                    raise json.JSONDecodeError("Debe existir pages[0] con width/height/background.", raw, 0)
                p0 = pages[0]
                if not isinstance(p0.get("width"), (int, float)) or not isinstance(p0.get("height"), (int, float)):
                    raise json.JSONDecodeError("width y height deben ser numéricos.", raw, 0)
                if not isinstance(p0.get("background"), str):
                    raise json.JSONDecodeError('background debe ser string (ej. "#ffffff").', raw, 0)
            except json.JSONDecodeError as e:
                return render(
                    request,
                    "administracion/plantilla_form.html",
                    {"mode": "create", "form_error": str(e), "prefill": {"title": title, "kind": kind, "json": raw}},
                )

        tpl = DesignTemplate.objects.create(title=title, kind=kind, json_active=data)
        DesignTemplateVersion.objects.create(
            template=tpl, version=1, json_payload=data, created_by=request.user, note="Versión inicial"
        )
        return redirect("administracion:plantilla_edit", tpl_id=tpl.id)

    return render(
    request,
    "administracion/plantilla_form.html",
    {
        "mode": "create",
        "prefill": {},
        # dict funciona con lookup por punto en Django templates
        "item": {"title": "", "kind": DesignTemplate.DESIGN, "json": ""},
    },
)

@login_required
@csp_update({
    'script-src': ("'self'", csp.NONCE),
    'style-src': ("'self'", "https://fonts.googleapis.com", csp.NONCE),
    'img-src': ("'self'", "data:", "blob:"),
})
@csp_update({'script-src': ("'self'", csp.NONCE, "https://cdn.jsdelivr.net")})
def plantilla_edit(request, tpl_id: int):
    tpl = get_object_or_404(DesignTemplate, pk=tpl_id)
    if request.method == "POST":
        try:
            data = json.loads(request.body.decode("utf-8"))
        except Exception:
            return HttpResponseBadRequest("Cuerpo inválido")
        tpl.json_active = data
        tpl.save(update_fields=["json_active", "updated_at"])
        return JsonResponse({"ok": True})

    assets = TemplateAsset.objects.all().order_by("-created_at")[:200]
    return render(
        request,
        "administracion/plantilla_editor.html",
        {"tpl": tpl, "assets": assets, "csrf_token": get_token(request)}
    )

@login_required
@require_http_methods(["POST"])
def plantilla_duplicate(request, tpl_id: int):
    tpl = get_object_or_404(DesignTemplate, pk=tpl_id)
    clone = DesignTemplate.objects.create(
        title=f"{tpl.title} (copia)",
        kind=tpl.kind,
        json_active=tpl.json_active,
        org=tpl.org
    )
    DesignTemplateVersion.objects.create(
        template=clone, version=1, json_payload=tpl.json_active, created_by=request.user, note="Copia"
    )
    return redirect("administracion:plantilla_edit", tpl_id=clone.id)

@login_required
@require_http_methods(["POST"])
def plantilla_new_version(request, tpl_id: int):
    tpl = get_object_or_404(DesignTemplate, pk=tpl_id)
    payload = request.POST.get("json")
    if not payload:
        return HttpResponseBadRequest("Falta json")
    try:
        data = json.loads(payload)
    except json.JSONDecodeError:
        return HttpResponseBadRequest("JSON inválido")

    last = DesignTemplateVersion.objects.filter(template=tpl).aggregate(m=Max("version")).get("m") or 0
    v = last + 1
    DesignTemplateVersion.objects.create(
        template=tpl, version=v, json_payload=data, created_by=request.user, note=request.POST.get("note", ""))
    tpl.json_active = data
    tpl.save(update_fields=["json_active", "updated_at"])
    return redirect("administracion:plantilla_edit", tpl_id=tpl.id)

@login_required
def assets_library(request):
    items = TemplateAsset.objects.all().order_by("-created_at")
    return render(request, "administracion/assets_library.html", {"items": items})

@login_required
@require_http_methods(["POST"])
def asset_upload(request):
    f = request.FILES.get("file")
    if not f:
        return HttpResponseBadRequest("Falta archivo")
    asset = TemplateAsset.objects.create(name=f.name, file=f, mime=f.content_type or "")
    return JsonResponse({"ok": True, "id": asset.id, "name": asset.name, "url": asset.file.url})

@login_required
@require_http_methods(["POST"])
def asset_delete(request, pk: int):
    """
    Elimina un asset y, si existe, el archivo físico en disco.
    Además intenta borrar la copia en static/img/diplomas o static/img/DC3
    generada por _mirror_png_to_static_background.
    """
    asset = get_object_or_404(TemplateAsset, pk=pk)

    # Guardamos datos ANTES de borrar el registro
    file_path  = asset.file.path if asset.file and hasattr(asset.file, "path") else None
    asset_name = asset.name

    # Borrar registro del modelo
    asset.delete()

    # 1) Borrar archivo en MEDIA
    if file_path:
        try:
            os.remove(file_path)
        except Exception:
            pass

    # 2) Borrar posible copia en static/img/...
    _delete_mirrored_static_by_asset_name(asset_name)

    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return JsonResponse({"ok": True})
    return redirect("administracion:assets_library")

# ========== PROGRAMAS / DIPLOMADOS ==========

@login_required
def program_create(request):
    plantillas = DesignTemplate.objects.all().order_by("title")

    # 👇 NUEVO: plantillas DOCX activas, separadas por tipo
    docx_tpls = DocxTemplate.objects.filter(is_active=True).order_by("name")
    docx_diploma_templates = docx_tpls.filter(tipo="diploma")
    docx_dc3_templates = docx_tpls.filter(tipo="dc3")
    docx_cproem_templates = docx_tpls.filter(tipo="cproem")

    if request.method == "POST":
        code = (request.POST.get("code") or "").strip().upper()
        name = (request.POST.get("name") or "").strip()
        raw_const = request.POST.get("constancia_type") or ConstanciaType.CEPROEM

        # Puede venir 0/1 o CEPROEM/DC3 o cproem/dc3
        if str(raw_const) in ("0", "1"):
            constancia = ConstanciaType.DC3 if str(raw_const) == "1" else ConstanciaType.CEPROEM
        else:
            constancia = raw_const

        plantilla_diploma_id = request.POST.get("plantilla_diploma") or None
        plantilla_constancia_id = request.POST.get("plantilla_constancia") or None

        # 👇 NUEVO: IDs de plantillas DOCX por tipo
        docx_tpl_diploma_id = request.POST.get("docx_tpl_diploma") or None
        docx_tpl_dc3_id = request.POST.get("docx_tpl_dc3") or None
        docx_tpl_cproem_id = request.POST.get("docx_tpl_cproem") or None

        item = {
            "code": code,
            "name": name,
            "constancia_type": constancia,
            "plantilla_diploma_id": plantilla_diploma_id,
            "plantilla_constancia_id": plantilla_constancia_id,
            # 👇 NUEVO: para que se mantengan seleccionados si hay error
            "docx_tpl_diploma_id": docx_tpl_diploma_id,
            "docx_tpl_dc3_id": docx_tpl_dc3_id,
            "docx_tpl_cproem_id": docx_tpl_cproem_id,
        }

        if not code or not name:
            messages.error(request, "Código y nombre son obligatorios.")
            return render(
                request,
                "administracion/program_form.html",
                {
                    "mode": "create",
                    "item": item,
                    "plantillas": plantillas,
                    "ConstanciaType": ConstanciaType,
                    "docx_diploma_templates": docx_diploma_templates,
                    "docx_dc3_templates": docx_dc3_templates,
                    "docx_cproem_templates": docx_cproem_templates,
                },
            )

        try:
            with transaction.atomic():
                # === 1) Crear Program en ces_db (tabla principal) ===
                program = Program.objects.create(
                    code=code,
                    name=name,
                    constancia_type=constancia,
                    plantilla_diploma_id=plantilla_diploma_id or None,
                    plantilla_constancia_id=plantilla_constancia_id or None,
                    # 👇 NUEVO: guardar FKs DOCX
                    docx_tpl_diploma_id=docx_tpl_diploma_id or None,
                    docx_tpl_dc3_id=docx_tpl_dc3_id or None,
                    docx_tpl_cproem_id=docx_tpl_cproem_id or None,
                )

                # === 2) Insertar en ces_simulacion.diplomado ===
                constancia_sim = 1 if constancia == ConstanciaType.DC3 else 0

                with connections["ces"].cursor() as cur:
                    cur.execute(
                        """
                        INSERT INTO diplomado (programa, programa_full, constancia, updated_at)
                        VALUES (%s, %s, %s, NOW())
                        """,
                        [code, name, constancia_sim],
                    )       

            messages.success(request, "Programa creado correctamente.")
            return redirect("administracion:program_list")

        except Exception as exc:
            messages.error(request, f"No se pudo crear el programa: {exc}")
            return render(
                request,
                "administracion/program_form.html",
                {
                    "mode": "create",
                    "item": item,
                    "plantillas": plantillas,
                    "ConstanciaType": ConstanciaType,
                    "docx_diploma_templates": docx_diploma_templates,
                    "docx_dc3_templates": docx_dc3_templates,
                    "docx_cproem_templates": docx_cproem_templates,
                },
            )

    # GET inicial: item vacío
    empty_item = {
        "code": "",
        "name": "",
        "constancia_type": ConstanciaType.CEPROEM,
        "plantilla_diploma_id": None,
        "plantilla_constancia_id": None,
        # 👇 NUEVO: sin DOCX por defecto
        "docx_tpl_diploma_id": None,
        "docx_tpl_dc3_id": None,
        "docx_tpl_cproem_id": None,
    }
    return render(
        request,
        "administracion/program_form.html",
        {
            "mode": "create",
            "item": empty_item,
            "plantillas": plantillas,
            "ConstanciaType": ConstanciaType,
            "docx_diploma_templates": docx_diploma_templates,
            "docx_dc3_templates": docx_dc3_templates,
            "docx_cproem_templates": docx_cproem_templates,
        },
    )


@csrf_exempt
@login_required
@require_http_methods(["GET", "POST"])
def program_edit(request, pk: int):
    """
    Edita el modelo Program (tabla Django) y sincroniza, en la medida de lo posible:
      - Program (BD principal)
      - alumnos.Program (ProgramSim)   [best effort]
      - ces_simulacion.diplomado       [best effort]

    El id que llega en la URL (pk) es el de Program.
    El id del diplomado (ces_simulacion.diplomado.id) viaja en sim_id (hidden del modal).
    """
    # Program de la BD principal
    p = get_object_or_404(Program, pk=pk)

    # ---------- RAMA AJAX (modal de edición) ----------
    if request.method == "POST" and request.headers.get("x-requested-with") == "XMLHttpRequest":
        # Parsear JSON
        try:
            data = json.loads(request.body.decode("utf-8") or "{}")
        except json.JSONDecodeError:
            return JsonResponse(
                {"success": False, "message": "Datos inválidos (JSON)."},
                status=400,
            )

        old_code = p.code

        new_code       = (data.get("programa") or "").strip()
        new_name       = (data.get("programa_full") or "").strip()
        new_constancia = (data.get("constancia") or "").strip().lower()

        # id del diplomado en ces_simulacion.diplomado
        sim_id_raw = data.get("sim_id")
        try:
            sim_id = int(sim_id_raw) if sim_id_raw not in (None, "",) else None
        except (TypeError, ValueError):
            sim_id = None

        # IDs de plantillas (pueden venir vacíos)
        def to_int_or_none(value):
            try:
                return int(value) if value not in (None, "",) else None
            except (TypeError, ValueError):
                return None

        diploma_design_id = to_int_or_none(data.get("diploma_design_id"))
        diploma_docx_id   = to_int_or_none(data.get("diploma_docx_id"))
        const_design_id   = to_int_or_none(data.get("const_design_id"))
        docx_dc3_id       = to_int_or_none(data.get("docx_dc3_id"))
        docx_cproem_id    = to_int_or_none(data.get("docx_cproem_id"))

        # ---- Validaciones básicas ----
        if not new_code or not new_name:
            return JsonResponse(
                {
                    "success": False,
                    "message": "Los campos «Programa» y «Nombre completo» son obligatorios.",
                },
                status=400,
            )

        if new_constancia not in ("dc3", "cproem"):
            return JsonResponse(
                {
                    "success": False,
                    "message": "Tipo de constancia inválido. Usa «dc3» o «cproem».",
                },
                status=400,
            )

        # Mapear dc3/cproem -> ConstanciaType y flag numérico (1/0) para ces_simulacion
        constancia_model = (
            ConstanciaType.DC3 if new_constancia == "dc3" else ConstanciaType.CEPROEM
        )
        constancia_flag = 1 if new_constancia == "dc3" else 0

        # ---- Validar unicidad en Program (para evitar IntegrityError) ----
        if Program.objects.exclude(pk=p.pk).filter(code=new_code).exists():
            return JsonResponse(
                {
                    "success": False,
                    "message": f"Ya existe otro programa con el código «{new_code}».",
                },
                status=400,
            )

        if Program.objects.exclude(pk=p.pk).filter(name=new_name).exists():
            return JsonResponse(
                {
                    "success": False,
                    "message": f"Ya existe otro programa con el nombre «{new_name}».",
                },
                status=400,
            )

        try:
            # 1) Guardar SIEMPRE el Program principal (crítico)
            with transaction.atomic():
                p.code = new_code
                p.name = new_name
                p.constancia_type = constancia_model

                # Plantillas (se asignan aunque vengan en blanco -> limpia el campo)
                # Usamos los campos *_id para no tener que buscar los objetos uno por uno.
                p.plantilla_diploma_id    = diploma_design_id
                p.docx_tpl_diploma_id     = diploma_docx_id
                p.plantilla_constancia_id = const_design_id
                p.docx_tpl_dc3_id         = docx_dc3_id
                p.docx_tpl_cproem_id      = docx_cproem_id

                p.save()

            # 2) Sincronizar alumnos.Program (ProgramSim) – best effort
            try:
                related_qs = None
                # No sabemos 100% el esquema real, así que usamos atributos defensivos
                if hasattr(ProgramSim, "abbreviation"):
                    related_qs = ProgramSim.objects.filter(abbreviation=old_code)
                elif hasattr(ProgramSim, "programa"):
                    related_qs = ProgramSim.objects.filter(programa=old_code)
                else:
                    related_qs = ProgramSim.objects.none()

                for pr in related_qs:
                    if hasattr(pr, "abbreviation"):
                        pr.abbreviation = new_code
                    if hasattr(pr, "programa"):
                        pr.programa = new_code
                    if hasattr(pr, "name"):
                        pr.name = new_name
                    if hasattr(pr, "programa_full"):
                        pr.programa_full = new_name
                    if hasattr(pr, "constancia_type"):
                        pr.constancia_type = constancia_model
                    pr.save()
            except Exception:
                logging.exception(
                    "No se pudo sincronizar alumnos.Program (ProgramSim); se continúa de todos modos."
                )

            # 3) Sincronizar ces_simulacion.diplomado – best effort
            try:
                with connections["ces"].cursor() as cur:
                    if sim_id is not None:
                        # Actualizamos por ID de diplomado (lo más seguro)
                        cur.execute(
                            """
                            UPDATE diplomado
                            SET programa      = %s,
                                programa_full = %s,
                                constancia    = %s,
                                updated_at    = NOW()
                            WHERE id = %s
                            """,
                            [new_code, new_name, constancia_flag, sim_id],
                        )
                    else:
                        # Fallback por código viejo, por si no vino el id
                        cur.execute(
                            """
                            UPDATE diplomado
                            SET programa      = %s,
                                programa_full = %s,
                                constancia    = %s,
                                updated_at    = NOW()
                            WHERE programa = %s
                            """,
                            [new_code, new_name, constancia_flag, old_code],
                        )
            except Exception:
                logging.exception(
                    "No se pudo sincronizar tabla diplomado en BD 'ces'; se continúa de todos modos."
                )

        except Exception as e:
            # Si algo falla al guardar Program, devolvemos error claro
            logging.exception("Error crítico al actualizar Program en program_edit")
            return JsonResponse(
                {
                    "success": False,
                    "message": f"No se pudo guardar el programa: {type(e).__name__}: {e}",
                },
                status=400,
            )

        # Si llegamos aquí, Program se guardó bien (y se intentó sincronizar todo)
        return JsonResponse(
            {
                "success": True,
                "message": f"El programa «{new_name}» se actualizó correctamente.",
                "data": {
                    "programa": new_code,
                    "programa_full": new_name,
                    "constancia": new_constancia,  # 'dc3' o 'cproem'
                },
            },
            status=200,
        )

    # ---------- RAMA NORMAL (formulario clásico) ----------
    if request.method == "POST":
        p.name = request.POST.get("name") or p.name
        p.code = request.POST.get("code") or p.code

        constancia_raw = (request.POST.get("constancia_type") or "").upper()
        if constancia_raw in (ConstanciaType.DC3, ConstanciaType.CEPROEM):
            p.constancia_type = constancia_raw

        plantilla_d = request.POST.get("plantilla_diploma") or None
        plantilla_c = request.POST.get("plantilla_constancia") or None

        p.plantilla_diploma = (
            DesignTemplate.objects.filter(id=plantilla_d).first()
            if plantilla_d
            else None
        )
        p.plantilla_constancia = (
            DesignTemplate.objects.filter(id=plantilla_c).first()
            if plantilla_c
            else None
        )

        # DOCX (si aún tienes esos campos en el form clásico)
        docx_d = request.POST.get("docx_tpl_diploma") or None
        docx_dc3 = request.POST.get("docx_tpl_dc3") or None
        docx_cproem = request.POST.get("docx_tpl_cproem") or None

        p.docx_tpl_diploma = (
            DocxTemplate.objects.filter(id=docx_d).first() if docx_d else None
        )
        p.docx_tpl_dc3 = (
            DocxTemplate.objects.filter(id=docx_dc3).first() if docx_dc3 else None
        )
        p.docx_tpl_cproem = (
            DocxTemplate.objects.filter(id=docx_cproem).first() if docx_cproem else None
        )

        p.save()
        return redirect("administracion:program_list")

    # ---------- GET normal: formulario de edición clásico ----------
    plantillas = DesignTemplate.objects.all().order_by("title")

    # Listas DOCX (por si las sigues usando en el form clásico)
    docx_tpls = DocxTemplate.objects.filter(is_active=True).order_by("name")
    docx_diploma_templates = docx_tpls.filter(tipo="diploma")
    docx_dc3_templates = docx_tpls.filter(tipo="dc3")
    docx_cproem_templates = docx_tpls.filter(tipo="cproem")

    return render(
        request,
        "administracion/program_form.html",
        {
            "mode": "edit",
            "item": p,
            "plantillas": plantillas,
            "ConstanciaType": ConstanciaType,
            "docx_diploma_templates": docx_diploma_templates,
            "docx_dc3_templates": docx_dc3_templates,
            "docx_cproem_templates": docx_cproem_templates,
        },
    )


@login_required
@require_http_methods(["POST"])
def plantilla_upload_thumb(request, tpl_id: int):
    """
    Sube una imagen manual y la guarda en DesignTemplate.thumb.
    No toca json_active ni el flujo de edición.
    """
    tpl = get_object_or_404(DesignTemplate, pk=tpl_id)
    f = request.FILES.get("thumb")
    if not f:
        return JsonResponse({"ok": False, "msg": "Falta archivo (thumb)."}, status=400)
    tpl.thumb = f
    tpl.save(update_fields=["thumb", "updated_at"])
    return JsonResponse({"ok": True, "url": tpl.thumb.url})

@login_required
@require_http_methods(["POST"])
def plantilla_regen_thumb(request, tpl_id: int):
    """
    Regenera la miniatura desde json_active y la guarda en thumb.
    Fuerza override para garantizar actualización.
    """
    tpl = get_object_or_404(DesignTemplate, pk=tpl_id)
    ok = save_thumb(tpl, force=True)
    if ok:
        tpl.save(update_fields=["thumb", "updated_at"])
        return JsonResponse({"ok": True, "url": tpl.thumb.url if tpl.thumb else None})
    return JsonResponse({"ok": False, "msg": "No se pudo generar la miniatura."}, status=500)

@login_required
@require_http_methods(["POST"])
def plantilla_delete(request, tpl_id: int):
    """
    Elimina una DesignTemplate y, además, intenta borrar la copia PNG en
    static/img/diplomas o static/img/DC3 que se generó al importar el documento.
    """
    tpl = get_object_or_404(DesignTemplate, pk=tpl_id)

    # Guardamos el título ANTES de borrar (ej. 'diploma-18.pdf', 'DC3__F.docx', etc.)
    title = tpl.title or ""

    # Borramos la plantilla
    tpl.delete()

    # Intentamos borrar el PNG espejado en static/img/...
    if title:
        _delete_mirrored_static_by_asset_name(title)

    messages.success(
        request,
        f"Se eliminó la plantilla «{title}»."
    )
    return redirect("administracion:plantillas_admin")

@login_required
@require_http_methods(["POST"])
def plantilla_import(request):
    """
    Importa un documento subido desde "Plantillas admin".

    Comportamiento:
    - Si es imagen (PNG/JPG/JPEG/WEBP/SVG) → crea una plantilla con la imagen como fondo.
    - Si es PDF u Office (DOC/DOCX/PPT/PPTX/ODT/ODP) → intenta generar PNG de la 1ª página
      usando `_pdf_to_png_bytes` / `_office_to_pdf_bytes`.
    - Si la conversión falla o no hay dependencias → crea una plantilla en blanco con
      un texto "Adjunto: <nombre>", pero el archivo original queda guardado como asset.
    - Siempre regresa al listado de plantillas (plantillas_admin); ya no abre el editor.
    """
    f = request.FILES.get("file")
    if not f:
        messages.error(request, "No se recibió ningún archivo.")
        return redirect("administracion:plantillas_admin")

    name = Path(f.name).name
    ext = Path(f.name).suffix.lower()

    # Guardamos SIEMPRE el archivo original como asset
    asset_original = TemplateAsset.objects.create(
        name=name,
        file=f,
        mime=f.content_type or "",
    )

    image_exts  = {".png", ".jpg", ".jpeg", ".webp", ".svg"}
    office_exts = {".doc", ".docx", ".ppt", ".pptx", ".odt", ".odp"}

    # 1) Imágenes → plantilla con imagen de fondo
    if ext in image_exts:
        data = {
            "pages": [{
                "width": 1920,
                "height": 1080,
                "background": "#ffffff",
                "layers": [{
                    "type": "image",
                    "url": asset_original.file.url,
                    "x": 0,
                    "y": 0,
                    "w": 1920,
                    "h": 1080,
                }],
            }]
        }

        tpl = DesignTemplate.objects.create(
            title=name,
            kind=DesignTemplate.DESIGN,
            json_active=data,
        )

        # Miniatura: usamos directamente la imagen original como thumb (si el modelo la tiene)
        try:
            if hasattr(tpl, "thumb") and asset_original.file:
                # leemos los bytes del archivo original
                with asset_original.file.open("rb") as fh:
                    img_bytes = fh.read()
                thumb_name = Path(name).name  # mismo nombre de archivo
                tpl.thumb.save(
                    thumb_name,
                    ContentFile(img_bytes, name=thumb_name),
                    save=True,
                )
        except Exception:
            # si algo falla, no rompemos el flujo
            pass

        messages.success(
            request,
            f"Se importó «{name}» como imagen de fondo. También se guardó el archivo original en la biblioteca de assets."
        )
        return redirect("administracion:plantillas_admin")

    # 2) PDF / Office → intentar generar PNG de 1ª página
    png_bytes = None
    try:
        if ext == ".pdf":
            with asset_original.file.open("rb") as fh:
                pdf_bytes = fh.read()
            png_bytes = _pdf_to_png_bytes(pdf_bytes)
        elif ext in office_exts:
            # Office → PDF usando helper ya definido más abajo
            with asset_original.file.open("rb") as fh:
                content = fh.read()

            class _U:
                def __init__(self, b): self._b = b
                def chunks(self, csize=64 * 1024):
                    mv = memoryview(self._b)
                    for i in range(0, len(mv), csize):
                        yield mv[i:i + csize]

            pdf_bytes = _office_to_pdf_bytes(_U(content), name_hint=name)
            if pdf_bytes:
                png_bytes = _pdf_to_png_bytes(pdf_bytes)
    except Exception:
        png_bytes = None

    # Si pudimos generar PNG → usarlo como fondo y como miniatura
    if png_bytes:
        img_name = Path(name).with_suffix(".png").name

        # 1) Asset en MEDIA (como ya hacías)
        asset_png = TemplateAsset.objects.create(
            name=img_name,
            file=ContentFile(png_bytes, name=img_name),
            mime="image/png",
        )

        # 2) Copia en static/img/diplomas o static/img/DC3
        _mirror_png_to_static_background(img_name, png_bytes)

        # 3) JSON activo con la imagen de fondo
        data = {
            "pages": [{
                "width": 1920,
                "height": 1080,
                "background": "#ffffff",
                "layers": [{
                    "type": "image",
                    "url": asset_png.file.url,
                    "x": 0,
                    "y": 0,
                    "w": 1920,
                    "h": 1080,
                }],
            }]
        }

        # 4) Crear la plantilla
        tpl = DesignTemplate.objects.create(
            title=name,
            kind=DesignTemplate.DESIGN,
            json_active=data,
        )

        # 5) Usar el mismo PNG como miniatura (thumb) para la tarjeta
        try:
            if hasattr(tpl, "thumb"):
                thumb_name = img_name  # por ejemplo "diploma-18.png"
                tpl.thumb.save(
                    thumb_name,
                    ContentFile(png_bytes, name=thumb_name),
                    save=True,
                )
        except Exception:
            pass

        messages.success(
            request,
            f"Se importó «{name}» y se generó una vista previa de la primera página."
        )
        return redirect("administracion:plantillas_admin")

    # 3) Fallback → plantilla en blanco con texto "Adjunto: <nombre>"
    data = {
        "pages": [{
            "width": 1280,
            "height": 720,
            "background": "#ffffff",
            "layers": [{
                "type": "text",
                "text": f"Adjunto: {name}",
                "x": 64,
                "y": 64,
                "fontSize": 32,
                "fill": "#111",
            }],
        }]
    }
    tpl = DesignTemplate.objects.create(
        title=name,
        kind=DesignTemplate.DESIGN,
        json_active=data,
    )

    messages.warning(
        request,
        "Se adjuntó el archivo original pero no se pudo generar la vista previa. "
        "Verifica que LibreOffice y Poppler estén instalados en el servidor si deseas convertir a imagen."
    )
    return redirect("administracion:plantillas_admin")


@login_required
@require_POST
def docx_template_import(request):
    f = request.FILES.get("file")
    if not f:
        messages.error(request, "No se recibió ningún archivo.")
        return redirect("administracion:plantillas_admin")

    ext = Path(f.name).suffix.lower()
    if ext not in {".doc", ".docx"}:
        messages.error(request, "Solo se permiten archivos .doc o .docx para plantillas con variables.")
        return redirect("administracion:plantillas_admin")

    tipo = (request.POST.get("tipo") or "diploma").lower()
    if tipo not in {"diploma", "dc3", "cproem"}:
        tipo = "diploma"

    titulo = Path(f.name).stem

    docx_tpl = DocxTemplate.objects.create(
        name=titulo,      # <--- campo correcto
        tipo=tipo,
        file=f,
        is_active=True,
    )

    messages.success(
        request,
        f"Se importó la plantilla DOCX «{docx_tpl.name}» para {tipo.upper()}."
    )
    return redirect("administracion:plantillas_admin")

@login_required
@require_http_methods(["POST"])
def docx_template_delete(request, pk: int):
    """
    Elimina una plantilla DOCX con variables.
    """
    tpl = get_object_or_404(DocxTemplate, pk=pk)
    nombre = tpl.name or "(sin nombre)"
    tpl.delete()
    messages.success(
        request,
        f"Se eliminó la plantilla DOCX «{nombre}».",
    )
    return redirect("administracion:plantillas_admin")

@login_required
@require_GET
def docx_template_preview(request, pk: int):
    """
    Devuelve una miniatura PNG de la primera página del DOCX.

    - Primera vez: DOCX -> PDF (LibreOffice) -> PNG (PyMuPDF) y se guarda
      como <nombre_docx>_preview.png en el mismo storage.
    - Siguientes veces: solo lee ese PNG ya generado (rápido).
    """
    tpl = get_object_or_404(DocxTemplate, pk=pk)

    if not tpl.file:
        return HttpResponseNotFound(b"No file")

    # -------- 1) Si ya existe la miniatura, la devolvemos directo --------
    base, _ = os.path.splitext(tpl.file.name)  # docx_templates/DAIEN_QR
    preview_relpath = f"{base}_preview.png"    # docx_templates/DAIEN_QR_preview.png

    if default_storage.exists(preview_relpath):
        with default_storage.open(preview_relpath, "rb") as fh:
            data = fh.read()
        return HttpResponse(data, content_type="image/png")

    # -------- 2) Generar miniatura (una sola vez) --------
    try:
        import fitz  # PyMuPDF
    except Exception:
        log.exception("PyMuPDF (fitz) no está disponible para generar vista previa DOCX.")
        return HttpResponseNotFound(b"preview-not-available")

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)

            # 2.1 Copiar DOCX a archivo temporal
            src_path = tmpdir_path / "src.docx"
            tpl.file.open("rb")
            try:
                src_path.write_bytes(tpl.file.read())
            finally:
                tpl.file.close()

            # 2.2 DOCX -> PDF con LibreOffice
            soffice = (
                getattr(settings, "SOFFICE_BIN", None)
                or os.environ.get("SOFFICE_BIN")
                or os.environ.get("LIBREOFFICE_PATH")
                or "soffice"
            )

            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--nologo",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmpdir_path),
                    str(src_path),
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )

            pdf_path = tmpdir_path / "src.pdf"
            if not pdf_path.exists():
                raise RuntimeError("No se generó el PDF temporal desde el DOCX.")

            # 2.3 PDF -> PNG (primera página) con un zoom moderado
            doc = fitz.open(pdf_path)
            try:
                page = doc.load_page(0)
                # 1.0x para que sea ligero (puedes subir a 1.3 si quieres más calidad)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.0, 1.0))
                png_bytes = pix.tobytes("png")
            finally:
                doc.close()

        # 2.4 Guardar PNG en el storage para reutilizarlo
        default_storage.save(preview_relpath, ContentFile(png_bytes))

        return HttpResponse(png_bytes, content_type="image/png")

    except Exception as e:
        log.exception("Error generando vista previa para DocxTemplate #%s: %s", tpl.id, e)
        return HttpResponseNotFound(b"preview-error")

def _pdf_to_png_bytes(pdf_bytes: bytes):
    """
    PDF -> PNG (1ª página) usando PyMuPDF (fitz).

    Ventajas:
    - No necesita Poppler ni binarios externos.
    - Todo se resuelve con la librería Python pymupdf, que ya está en requirements.txt.
    """
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if doc.page_count == 0:
            return None

        page = doc.load_page(0)          # primera página (índice 0)
        zoom = 2.0                       # 2x para que se vea nítido
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)

        # Devolver bytes PNG directamente
        png_bytes = pix.tobytes("png")
        return png_bytes
    except Exception:
        return None

def _office_to_pdf_bytes(uploaded_file, name_hint: str):
    """DOC/DOCX/PPT/PPTX/ODT/ODP -> PDF con LibreOffice (headless)."""
    soffice = _find_soffice()
    if not soffice:
        return None
    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / name_hint
        with open(src, "wb") as out:
            for ch in uploaded_file.chunks():
                out.write(ch)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", td, str(src)],
                check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
            )
        except Exception:
            return None
        pdf_path = src.with_suffix(".pdf")
        if not pdf_path.exists():
            cand = list(Path(td).glob("*.pdf"))
            pdf_path = cand[0] if cand else None
        return pdf_path.read_bytes() if pdf_path and pdf_path.exists() else None
# ============= NUEVA VISTA =============
@login_required
@require_http_methods(["POST"])
def import_office_to_image(request):
    """
    Endpoint AJAX: recibe 'file', convierte a imagen (1a página),
    crea una DesignTemplate con esa imagen como fondo y responde JSON.
    """
    up = request.FILES.get("file")
    if not up:
        return JsonResponse({"ok": False, "error": "Falta archivo."}, status=400)

    name = Path(up.name).name
    ext  = Path(up.name).suffix.lower()

    png_bytes = None
    if ext == ".pdf":
        pdf_bytes = b"".join(up.chunks())
        png_bytes = _pdf_to_png_bytes(pdf_bytes)
    elif ext in {".doc", ".docx", ".ppt", ".pptx", ".odt", ".odp"}:
        pdf_bytes = _office_to_pdf_bytes(up, name)
        if pdf_bytes:
            png_bytes = _pdf_to_png_bytes(pdf_bytes)
    elif ext in {".png", ".jpg", ".jpeg", ".webp", ".svg"}:
        # si es imagen, simplemente la guardamos
        from .models import TemplateAsset, DesignTemplate
        asset = TemplateAsset.objects.create(name=name, file=up, mime=up.content_type or "")
        data = {
            "pages": [{
                "width": 1920, "height": 1080, "background": "#ffffff",
                "layers": [{"type":"image","url":asset.file.url,"x":0,"y":0,"w":1920,"h":1080}]
            }]
        }
        tpl = DesignTemplate.objects.create(title=name, kind=DesignTemplate.DESIGN, json_active=data)
        return JsonResponse({"ok": True, "template_id": tpl.id, "redirect": reverse("administracion:plantilla_edit", args=[tpl.id])})

    if not png_bytes:
        return JsonResponse({
            "ok": False,
            "error": "No se pudo convertir el documento. Instala Poppler (pdf2image) y LibreOffice para Office."
        }, status=501)

    # guardar PNG como asset y crear plantilla
    from .models import TemplateAsset, DesignTemplate
    img_name = Path(name).with_suffix(".png").name
    asset = TemplateAsset.objects.create(
        name=img_name,
        file=ContentFile(png_bytes, name=img_name),
        mime="image/png",
    )
    data = {
        "pages": [{
            "width": 1920, "height": 1080, "background": "#ffffff",
            "layers": [{"type":"image","url":asset.file.url,"x":0,"y":0,"w":1920,"h":1080}]
        }]
    }
    tpl = DesignTemplate.objects.create(title=name, kind=DesignTemplate.DESIGN, json_active=data)
    return JsonResponse({"ok": True, "template_id": tpl.id, "redirect": reverse("administracion:plantilla_edit", args=[tpl.id])})

def _find_soffice():
    # Si definiste SOFFICE_BIN, úsalo
    if SOFFICE_BIN and Path(SOFFICE_BIN).exists():
        return SOFFICE_BIN
    # Rutas típicas en Windows
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice",  # si está en PATH
    ]
    for c in candidates:
        p = Path(c)
        if (p.exists() and p.is_file()) or c == "soffice":
            return c
    return None

@login_required
@require_http_methods(["POST"])
def asset_convert_to_image(request, asset_id: int):
    """
    Convierte un asset (PDF/Office/imagen) a PNG (1ª página) y devuelve la URL del nuevo asset PNG.
    No borra el original.
    """
    asset = get_object_or_404(TemplateAsset, pk=asset_id)
    name = Path(asset.name).name
    ext  = Path(asset.name).suffix.lower()

    png_bytes = None
    if ext == ".pdf":
        # leer bytes del archivo original
        with asset.file.open("rb") as fh:
            pdf_bytes = fh.read()
        png_bytes = _pdf_to_png_bytes(pdf_bytes)
    elif ext in {".doc", ".docx", ".ppt", ".pptx", ".odt", ".odp"}:
        with asset.file.open("rb") as fh:
            # fake UploadedFile para reuse de helper
            class _U:
                def __init__(self, b): self._b = b
                def chunks(self, csize=64*1024):
                    yield self._b
            pdf_bytes = _office_to_pdf_bytes(_U(fh.read()), name)
        if pdf_bytes:
            png_bytes = _pdf_to_png_bytes(pdf_bytes)
    elif ext in {".png", ".jpg", ".jpeg", ".webp"}:
        with asset.file.open("rb") as fh:
            png_bytes = fh.read() if ext == ".png" else None  # ya es imagen (si jpg/webp podrías convertir)
        if not png_bytes:
            try:
                from PIL import Image
                im = Image.open(asset.file)
                if im.mode != "RGB":
                    im = im.convert("RGB")
                buff = io.BytesIO()
                im.save(buff, format="PNG")
                png_bytes = buff.getvalue()
            except Exception:
                png_bytes = None
    else:
        return JsonResponse({"ok": False, "error": "Formato no soportado para convertir."}, status=400)

    if not png_bytes:
        return JsonResponse({
            "ok": False,
            "error": "No se pudo convertir. Revisa Poppler (pdf2image) y, si es Office, LibreOffice (soffice)."
        }, status=501)

    # Guardar nuevo asset PNG
    img_name = Path(name).with_suffix(".png").name
    new_asset = TemplateAsset.objects.create(
        name=img_name,
        file=ContentFile(png_bytes, name=img_name),
        mime="image/png",
    )

    # Copia en static/img/diplomas o static/img/DC3
    _mirror_png_to_static_background(img_name, png_bytes)

    return JsonResponse({
        "ok": True,
        "id": new_asset.id,
        "name": new_asset.name,
        "url": new_asset.file.url,
    })

@require_POST
@login_required
@csrf_protect
def program_bulk_delete(request):
    # Espera JSON: { "ids": [1,2,3,...] }
    import json
    try:
        data = json.loads(request.body.decode("utf-8"))
        ids = data.get("ids", [])
    except Exception:
        return JsonResponse({"ok": False, "error": "JSON inválido."}, status=400)

    if not isinstance(ids, list):
        return JsonResponse({"ok": False, "error": "ids debe ser una lista."}, status=400)

    # Autorización adicional si aplica…
    Program.objects.filter(id__in=ids).delete()
    return JsonResponse({"ok": True, "deleted": len(ids)})

def _query_simulacion(sql, params=None):
    with connections['ces'].cursor() as cur:
        cur.execute(sql, params or [])
        # Si es SELECT habrá description; si no, es escritura.
        if cur.description:
            cols = [c[0] for c in cur.description]
            return [dict(zip(cols, row)) for row in cur.fetchall()]
        return cur.rowcount  # para DELETE/UPDATE/INSERT

def _query_default(sql, params=None):
    with connections['default'].cursor() as cur:
        cur.execute(sql, params or [])
        if cur.description:
            cols = [c[0] for c in cur.description]
            return [dict(zip(cols, row)) for row in cur.fetchall()]
        return cur.rowcount

@login_required
@require_POST
def program_delete(request, pk: int):
    """
    Elimina un diplomado de la BD de simulación (tabla `diplomado` en conexión 'ces'),
    pero SOLO si no hay alumnos con solicitudes activas para ese programa.
    """

    # 1) Buscar el diplomado en la BD de simulación y obtener sus datos
    with connections["ces"].cursor() as cur:
        cur.execute(
            "SELECT id, programa, programa_full FROM diplomado WHERE id = %s",
            [pk],
        )
        row = cur.fetchone()

    if not row:
        messages.error(request, "El diplomado que intentas borrar no existe.")
        return redirect("administracion:program_list")

    dip_id, abbreviation, full_name = row

    # 2) Localizar el/los programas correspondientes en la BD real (alumnos_program)
    related_programs = ProgramSim.objects.filter(abbreviation=abbreviation)

    # Estados que BLOQUEAN el borrado (inglés/español como comentaste)
    BLOCKED_STATUSES = [
        "pending", "pendiente",
        "review", "revision",
        "accepted", "aprobada",
        "rejected", "rechazada",
        "generating", "generando",
        "emailed", "enviada",
        "downloaded", "descargada", "download",
    ]

    blocked_qs = Request.objects.none()
    if related_programs.exists():
        blocked_qs = Request.objects.filter(
            program__in=related_programs,
            status__in=BLOCKED_STATUSES,
        )

    blocked_count = blocked_qs.count()

    # 3) Si hay alumnos con solicitudes en esos estatus → NO se borra
    if blocked_count > 0:
        messages.error(
            request,
            (
                f"No se puede eliminar el diplomado «{full_name}» porque hay "
                f"{blocked_count} alumno(s) con solicitudes en proceso o ya generadas. "
                "Solo podrás eliminarlo cuando ya no haya alumnos cursando ni con "
                "solicitudes pendientes; es decir, cuando todas estén en estatus "
                "«completado»."
            ),
        )
        return redirect("administracion:program_list")

    # 4) Si no hay solicitudes bloqueantes, borrar el diplomado en la simulación
    with transaction.atomic(using="ces"):
        with connections["ces"].cursor() as cur:
            cur.execute("DELETE FROM diplomado WHERE id = %s", [dip_id])

    messages.success(
        request,
        f"El diplomado «{full_name}» se eliminó correctamente."
    )
    return redirect("administracion:program_list")

# ---------- LÓGICA DE NEGOCIO ----------
def fetch_programs(search_text=None):
    """
    Une programas de:
      - ces_simulacion.diplomado
      - ces_db.alumnos_program
    Estandariza campos con alias: codigo, nombre, modalidad, fuente.
    """
    like = f"%{search_text}%" if search_text else None

    # OJO: usa la PK real de cada tabla y dale alias AS codigo
    # Ajusta los nombres 'nombre' y 'modalidad' si en tu esquema se llaman distinto
    sql_sim = """
        SELECT
            d.id            AS codigo,     -- si tu PK es otra (p.ej. id_diplomado/clave), cámbiala aquí
            d.nombre        AS nombre,     -- cámbialo si tu columna es 'titulo' u otro
            d.modalidad     AS modalidad,  -- si no existe, puedes devolver NULL
            'simulacion'    AS fuente
        FROM diplomado d
        WHERE (%s IS NULL OR d.nombre LIKE %s)
        ORDER BY d.id DESC
    """
    sim_params = [like, like]

    sql_def = """
        SELECT
            p.id            AS codigo,     -- si tu tabla tiene 'codigo' úsala; si no, deja 'id'
            p.nombre        AS nombre,
            p.modalidad     AS modalidad,
            'alumnos_program' AS fuente
        FROM alumnos_program p
        WHERE (%s IS NULL OR p.nombre LIKE %s)
        ORDER BY p.id DESC
    """
    def_params = [like, like]

    sim_rows = _query_simulacion(sql_sim, sim_params)
    def_rows = _query_default(sql_def, def_params)

    # Unimos resultados
    items = sim_rows + def_rows
    return items
def _table_columns(using_alias: str, schema: str, table: str) -> set[str]:
    """
    Devuelve el conjunto de columnas existentes para schema.table en la conexión 'using_alias'.
    using_alias típicos: 'ces' (ces_simulacion), 'default' (ces_db).
    """
    with connections[using_alias].cursor() as cur:
        cur.execute(
            """
            SELECT COLUMN_NAME
            FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_SCHEMA=%s AND TABLE_NAME=%s
            """,
            [schema, table],
        )
        return {row[0] for row in cur.fetchall()}

def _pick(cols: set[str], candidates: list[str]) -> str | None:
    """Devuelve el primer nombre de columna candidato que exista en 'cols'."""
    for c in candidates:
        if c in cols:
            return c
    return None

def _safe_like(q: str | None) -> str | None:
    return f"%{q.strip()}%" if q else None

def _build_select(schema: str, table: str, alias: str, cols: set[str], fuente: str, q: str | None):
    """
    Construye un SELECT que normaliza:
      - codigo (PK o ‘clave’)
      - nombre  (título del programa)
      - modalidad (si existe; si no, NULL)
      - fuente  (constante para saber de dónde viene)
    Retorna: (sql, params) o (None, None) si no hay al menos codigo+nombre.
    """
    # Posibles nombres en tus tablas
    code_opts      = ["codigo", "clave", "cod", "cve", "id_diplomado", "id", "clave_diplomado", "codigo_prog", "cod_programa"]
    name_opts      = ["nombre", "nombre_diplomado", "nom_diplomado", "programa", "titulo", "descripcion", "nombre_programa", "name"]
    modality_opts  = ["modalidad", "modalidad_nombre", "tipo", "tipo_modalidad", "modalidad_prog"]

    col_code = _pick(cols, code_opts)
    col_name = _pick(cols, name_opts)
    col_mod  = _pick(cols, modality_opts)  # puede no existir

    if not col_code or not col_name:
        return None, None

    select_parts = [
        f"{alias}.`{col_code}` AS codigo",
        f"{alias}.`{col_name}` AS nombre",
        f"{alias}.`{col_mod}` AS modalidad" if col_mod else "NULL AS modalidad",
        "%s AS fuente",
    ]

    where_parts = []
    params = [fuente]

    like_q = _safe_like(q)
    if like_q:
        # filtra por nombre o por código textual si aplica
        where_parts.append(f"({alias}.`{col_name}` LIKE %s OR CAST({alias}.`{col_code}` AS CHAR) LIKE %s)")
        params.extend([like_q, like_q])

    sql = f"SELECT {', '.join(select_parts)} FROM `{schema}`.`{table}` {alias}"
    if where_parts:
        sql += " WHERE " + " AND ".join(where_parts)

    return sql, params

def _query_dicts(using_alias: str, sql: str, params: list) -> list[dict]:
    """Ejecuta y devuelve lista de dicts con nombres de columna."""
    with connections[using_alias].cursor() as cur:
        cur.execute(sql, params)
        rows = cur.fetchall()
        names = [d[0] for d in cur.description]
    return [dict(zip(names, r)) for r in rows]

def fetch_programs(search_text: str | None = None) -> list[dict]:
    """
    Une programas desde:
      - ces_simulacion.diplomado  (conexión 'ces')     → fuente = 'simulacion'
      - ces_db.alumnos_program     (conexión 'default')→ fuente = 'alumnos_program'
    Descubre columnas reales y normaliza a: codigo, nombre, modalidad, fuente.
    """
    items: list[dict] = []

    # ---------- SIMULACION ----------
    try:
        ces_schema = settings.DATABASES["ces"]["NAME"]      # usualmente 'ces_simulacion'
    except Exception:
        ces_schema = "ces_simulacion"
    sim_table  = "diplomado"
    sim_cols   = _table_columns("ces", ces_schema, sim_table)

    sim_sql, sim_params = _build_select(
        schema=ces_schema, table=sim_table, alias="d",
        cols=sim_cols, fuente="simulacion", q=search_text
    )
    if sim_sql:
        try:
            items.extend(_query_dicts("ces", sim_sql, sim_params))
        except Exception:
            # si falla esta fuente, seguimos con la otra
            pass

    # ---------- APP (DEFAULT) ----------
    try:
        def_schema = settings.DATABASES["default"]["NAME"]  # usualmente 'ces_db'
    except Exception:
        def_schema = "ces_db"
    app_table = "alumnos_program"
    app_cols  = _table_columns("default", def_schema, app_table)

    app_sql, app_params = _build_select(
        schema=def_schema, table=app_table, alias="p",
        cols=app_cols, fuente="alumnos_program", q=search_text
    )
    if app_sql:
        try:
            items.extend(_query_dicts("default", app_sql, app_params))
        except Exception:
            pass

    # Ordena por nombre y luego por código para vista estable
    items.sort(key=lambda x: ((x.get("nombre") or "").lower(), str(x.get("codigo") or "")))
    return items

def _dictfetchall(cursor):
    """Convierte resultados de cursor en lista de dicts."""
    cols = [col[0] for col in cursor.description]
    return [dict(zip(cols, row)) for row in cursor.fetchall()]

@csrf_exempt  # durante desarrollo te quita dolores de cabeza con CSRF
@login_required
@require_http_methods(["POST"])
def program_edit_api(request, pk: int):
    try:
        data = json.loads(request.body.decode("utf-8") or "{}")
    except json.JSONDecodeError:
        return JsonResponse(
            {"success": False, "message": "Datos inválidos."},
            status=400,
        )

    prog_admin = get_object_or_404(ProgramAdmin, pk=pk)
    old_code = prog_admin.code

    new_code = (data.get("programa") or "").strip()
    new_name = (data.get("programa_full") or "").strip()
    new_constancia = (data.get("constancia") or "").strip().lower()

    if not new_code or not new_name:
        return JsonResponse(
            {
                "success": False,
                "message": "Los campos «Programa» y «Nombre completo» son obligatorios.",
            },
            status=400,
        )

    if new_constancia not in ("dc3", "cproem"):
        return JsonResponse(
            {
                "success": False,
                "message": "Tipo de constancia inválido. Usa «dc3» o «cproem».",
            },
            status=400,
        )

    constancia_flag = 1 if new_constancia == "dc3" else 0

    try:
        with transaction.atomic():
            # 1) BD admin
            prog_admin.code = new_code
            prog_admin.name = new_name
            prog_admin.constancia_type = new_constancia
            prog_admin.save()

            # 2) BD alumnos_program (ProgramSim)
            related = ProgramSim.objects.filter(abbreviation=old_code)
            for p in related:
                p.abbreviation = new_code
                if hasattr(p, "name"):
                    p.name = new_name
                if hasattr(p, "constancia_type"):
                    p.constancia_type = new_constancia
                p.save()

            # 3) BD ces_simulacion.diplomado
            with connections["ces"].cursor() as cur:
                cur.execute(
                    """
                    UPDATE diplomado
                    SET programa = %s,
                        programa_full = %s,
                        constancia = %s,
                        update_at = NOW()
                    WHERE programa = %s
                    """,
                    [new_code, new_name, constancia_flag, old_code],
                )
    except Exception:
        return JsonResponse(
            {
                "success": False,
                "message": (
                    "Ocurrió un error al guardar los cambios. "
                    "Intenta de nuevo o contacta al administrador."
                ),
            },
            status=500,
        )

    return JsonResponse(
        {
            "success": True,
            "message": f"El programa «{new_name}» se actualizó correctamente.",
            "data": {
                "programa": new_code,
                "programa_full": new_name,
                "constancia": new_constancia,
            },
        }
    )
@login_required
def program_list(request):
    """
    Lista de programas de ces_simulacion (diplomado) y asegura
    que exista un Program para cada código de programa.
    Maneja duplicados de name de forma segura.
    """
    q = (request.GET.get("q") or "").strip()

    # ============================
    # PLANTILLAS (DC3, CPROEM y DIPLOMA)
    # ============================
    # Reutilizamos el mismo queryset para las 3 listas del modal (HTML/imagen)
    all_templates = DesignTemplate.objects.filter(kind="design").order_by("title")
    dc3_templates = all_templates
    cproem_templates = all_templates
    diploma_templates = all_templates   # <-- para el select de diploma

    # Plantillas DOCX activas (para el modal de edición)
    docx_tpls = DocxTemplate.objects.filter(is_active=True).order_by("name")
    docx_diploma_templates = docx_tpls.filter(tipo="diploma")
    docx_dc3_templates = docx_tpls.filter(tipo="dc3")
    docx_cproem_templates = docx_tpls.filter(tipo="cproem")



    # Program (tabla nueva) para la lista de la derecha (si la usas)
    programs_qs = Program.objects.all()
    if q:
        programs_qs = programs_qs.filter(
            Q(code__icontains=q) | Q(name__icontains=q)
        )
    programs = programs_qs.order_by("id")

    # Programas de ces_simulacion.diplomado
    with connections["ces"].cursor() as cur:
        sql = """
            SELECT d.id, d.programa, d.programa_full, d.constancia
            FROM diplomado AS d
        """
        params = []
        if q:
            sql += " WHERE d.programa LIKE %s OR d.programa_full LIKE %s"
            like = f"%{q}%"
            params.extend([like, like])
        sql += " ORDER BY d.id"

        cur.execute(sql, params)
        rows = cur.fetchall()

    sim_programs = []

    for row in rows:
        diplomado_id, programa, programa_full, constancia_flag = row
        if not programa:
            continue

        constancia_ui = "dc3" if constancia_flag == 1 else "cproem"
        constancia_model = (
            ConstanciaType.DC3 if constancia_flag == 1 else ConstanciaType.CEPROEM
        )

        # --- Asegurar Program para este código, manejando duplicados de name ---
        prog_obj = None

        # 1) Intentar por code primero (lo normal)
        prog_obj = Program.objects.filter(code=programa).first()

        if not prog_obj:
            # 2) No existe por code; intentamos crearlo
            base_name = programa_full or programa or f"Programa {diplomado_id}"

            try:
                prog_obj = Program.objects.create(
                    code=programa,
                    name=base_name,
                    constancia_type=constancia_model,
                )
            except IntegrityError:
                # 3) El name ya existe para otro registro -> buscamos uno existente
                #    o inventamos un nombre alterno único.
                # 3a) Si existe por name, lo reutilizamos
                prog_obj = Program.objects.filter(name=base_name).first()

                if not prog_obj:
                    # 3b) Generar nombre único sin romper unique(name)
                    safe_name = base_name
                    suffix = 2
                    while Program.objects.filter(name=safe_name).exists():
                        safe_name = f"{base_name} ({suffix})"
                        suffix += 1

                    prog_obj = Program.objects.create(
                        code=programa,
                        name=safe_name,
                        constancia_type=constancia_model,
                    )

                # ============================
        # Plantillas guardadas en Program
        # ============================
        tpl_constancia_id = None
        tpl_diploma_id = None
        docx_diploma_id = None
        docx_dc3_id = None
        docx_cproem_id = None

        if prog_obj:
            tpl_constancia_id = prog_obj.plantilla_constancia_id
            tpl_diploma_id = prog_obj.plantilla_diploma_id  # fondo diploma

            # DOCX
            docx_diploma_id = prog_obj.docx_tpl_diploma_id
            docx_dc3_id = prog_obj.docx_tpl_dc3_id
            docx_cproem_id = prog_obj.docx_tpl_cproem_id


        # Para DC3: frontal = plantilla_constancia (por ahora), reverso vacío
        if constancia_ui == "dc3":
            tpl_dc3_front_id = tpl_constancia_id
            tpl_dc3_back_id = None
            tpl_cproem_id = None
        else:
            # Para CPROEM usamos plantilla_constancia como plantilla CPROEM
            tpl_dc3_front_id = None
            tpl_dc3_back_id = None
            tpl_cproem_id = tpl_constancia_id

        # Llenamos la lista para la tabla de la izquierda (ces_simulacion)
        sim_programs.append(
            {
                "id": diplomado_id,
                "programa": programa,
                "programa_full": programa_full,
                "constancia": constancia_ui,
                "program_admin_id": prog_obj.id if prog_obj else None,
                "tpl_dc3_front_id": tpl_dc3_front_id,
                "tpl_dc3_back_id": tpl_dc3_back_id,
                "tpl_cproem_id": tpl_cproem_id,
                "tpl_diploma_id": tpl_diploma_id,
                # NUEVOS: para el modal de edición
                "tpl_constancia_id": tpl_constancia_id,
                "docx_diploma_id": docx_diploma_id,
                "docx_dc3_id": docx_dc3_id,
                "docx_cproem_id": docx_cproem_id,
            }
        )

    return render(
        request,
        "administracion/program_list.html",
                {
            "programs": programs,
            "sim_programs": sim_programs,
            "q": q,
            "dc3_templates": dc3_templates,
            "cproem_templates": cproem_templates,
            "diploma_templates": diploma_templates,  # HTML/imagen
            "docx_diploma_templates": docx_diploma_templates,
            "docx_dc3_templates": docx_dc3_templates,
            "docx_cproem_templates": docx_cproem_templates,
        },
    )



def verificar_documento(request, token):
    # 1) Buscar el token activo
    tok = get_object_or_404(
        DocToken.objects.select_related("request__program"),
        token=token,
        is_active=True,
    )

    req = tok.request          # alumnos.Request
    prog = getattr(req, "program", None)

    # 2) Datos base
    tipo = tok.tipo  # "diploma", "dc3", "cproem"
    alumno = f"{req.name} {req.lastname}".strip()
    programa = prog.name if prog else "Programa no disponible"

    # Graduate enlazado a Request (OneToOne)
    grad = getattr(req, "graduate", None)

    if grad:
        inicio = grad.validity_start
        fin = grad.validity_end
        folio = grad.curp or f"GRAD-{grad.id}"
    else:
        inicio = None
        fin = None
        folio = f"REQ-{req.id}"

    # ⏰ Momento exacto de la verificación
    verified_at = timezone.localtime()

    ctx = {
        "tipo": tipo,
        "alumno": alumno,
        "programa": programa,
        "inicio": inicio,
        "fin": fin,
        "folio": folio,
        "verified_at": verified_at,
        # extra: por si en algún lado sigues usando {{ now }}
        "now": verified_at,
    }

    return render(request, "administracion/verify_result.html", ctx)

def _mirror_png_to_static_background(img_name: str, png_bytes: bytes) -> str | None:
    """
    Guarda una copia del PNG en:
      - static/img/DC3/<NOMBRE>.png  si el nombre contiene 'DC3'
      - static/img/diplomas/<NOMBRE>.png  para el resto

    Ejemplos:
      DC3_DAP_F.pdf      -> static/img/DC3/DC3_DAP_F.png
      DC3__R.png         -> static/img/DC3/DC3__R.png
      Diploma_CPROEM_WEB -> static/img/diplomas/DIPLOMA_CPROEM_WEB.png
      Constancia_DAP     -> static/img/diplomas/CONSTANCIA_DAP.png
    """
    try:
        base = Path(img_name).stem.upper()  # nombre sin extensión

        # Regla DC3: cualquier nombre que contenga 'DC3'
        is_dc3 = "DC3" in base

        if is_dc3:
            subdir = "DC3"
        else:
            subdir = "diplomas"

        filename   = f"{base}.png"
        static_dir = Path(settings.BASE_DIR) / "static" / "img" / subdir
        static_dir.mkdir(parents=True, exist_ok=True)

        target = static_dir / filename
        with open(target, "wb") as fh:
            fh.write(png_bytes)

        return str(target)
    except Exception:
        # best effort; si falla no rompemos el flujo
        return None
    
@login_required
@require_http_methods(["POST"])
def plantilla_import_convert(request):
    # Simplemente reutiliza la lógica de plantilla_import
    return plantilla_import(request)

def _delete_mirrored_static_by_asset_name(asset_name: str) -> None:
    """
    Borra la copia PNG en static/img/diplomas o static/img/DC3
    calculando el nombre EXACTAMENTE igual que en _mirror_png_to_static_background.
    Acepta tanto 'diploma-18.pdf' como 'diploma-18.png'.
    """
    try:
        stem = Path(asset_name).stem          # 'diploma-18', 'DC3__F', etc.
        base_upper = stem.upper()             # 'DIPLOMA-18', 'DC3__F'

        # Misma regla que en _mirror_png_to_static_background:
        if base_upper.startswith("DC3_"):
            subdir = "DC3"
        else:
            subdir = "diplomas"

        filename   = f"{base_upper}.png"      # DIPLOMA-18.png, DC3__F.png
        static_dir = Path(settings.BASE_DIR) / "static" / "img" / subdir
        target     = static_dir / filename

        if target.exists():
            try:
                os.remove(target)
            except Exception:
                pass
    except Exception:
        # best-effort, nunca romper el flujo
        pass
    
def _absolute_static(request, path):
    """
    Devuelve la URL ABSOLUTA para un archivo estático,
    ej. http://127.0.0.1:8000/static/...
    """
    return request.build_absolute_uri(static(path))

def _get_program_from_graduate(grad):
    """
    Intenta obtener el Program asociado al Graduate,
    pasando por Request si existe esa relación.
    """
    req = getattr(grad, "request", None)
    if req is not None:
        return getattr(req, "program", None)
    return None


def pdf_cproem_digital(request, graduate_id, *args, **kwargs):
    """
    Genera el CPROEM digital del egresado con id = graduate_id.
    URL esperada: /administracion/cproem/digital/<int:graduate_id>/
    """
    graduate = get_object_or_404(Graduate, pk=graduate_id)

    # Contexto completo CPROEM
    context = _build_cproem_context(graduate, request)

    # Template DIGITAL
    html = render_to_string(
        "administracion/pdf_constancia_cproem_digital.html",
        context,
    )

    pdf_bytes = HTML(
        string=html,
        base_url=request.build_absolute_uri("/"),
    ).write_pdf()

    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    response["Content-Disposition"] = (
        f'inline; filename="constancia_cproem_digital_{graduate_id}.pdf"'
    )
    return response

def pdf_cproem_impreso(request, graduate_id: int):
    """
    Versión IMPRESA de la constancia CPROEM.
    Usa el mismo contexto que el CPROEM digital.
    """
    graduate = get_object_or_404(Graduate, pk=graduate_id)

    # Usamos el contexto "gordo" que ya armamos para el digital
    context = _build_cproem_context(graduate, request)

    html = render_to_string(
        "administracion/pdf_constancia_cproem_impresa.html",
        context,
    )

    pdf = HTML(
        string=html,
        base_url=request.build_absolute_uri("/"),
    ).write_pdf()

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = (
        f'inline; filename="constancia_cproem_impresa_{graduate_id}.pdf"'
    )
    return response

@login_required
@require_POST
def request_delete_rejected(request, pk):
    req = Request.objects.filter(id=pk).first()
    if not req:
        messages.error(request, "Solicitud no encontrada.")
        return redirect("administracion:requests_board")

    # Guardar en archivo histórico antes de borrar
    RejectedArchive.objects.create(
        email=req.email,
        full_name=f"{req.name} {req.lastname}",
        reason=req.status_reason or "Solicitud rechazada permanentemente.",
    )

    # Eliminar la solicitud real
    req.delete()

    messages.success(request, "Solicitud rechazada archivada y eliminada correctamente.")
    return redirect("administracion:requests_board")


def requests_board(request):
    estado = request.GET.get("estado") or "pending"
    q = request.GET.get("q", "").strip()
    program_id = request.GET.get("program")

    qs = Request.objects.all()

    # 🔹 NO mostrar las eliminadas
    qs = qs.filter(is_deleted=False)

    if estado:
        qs = qs.filter(status=estado)

    if q:
        qs = qs.filter(
            Q(email__icontains=q) |
            Q(name__icontains=q) |
            Q(lastname__icontains=q)
        )

    if program_id:
        qs = qs.filter(program_id=program_id)

    items = qs.order_by("-sent_at")
    # ... resto igual ...
